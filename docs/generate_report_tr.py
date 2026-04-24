"""
Online Hastane Yönetim Sistemi için profesyonel bir Word Belgesi (.docx) raporu oluşturur.
Akademik sunum için uygundur - Türkçe versiyon.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# Sayfa Düzeni
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Stil Yapılandırması
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for level in range(1, 4):
    heading_style = doc.styles['Heading %d' % level]
    heading_style.font.color.rgb = RGBColor(0x0d, 0x47, 0xa1)
    heading_style.font.bold = True
    if level == 1:
        heading_style.font.size = Pt(20)
    elif level == 2:
        heading_style.font.size = Pt(15)
    else:
        heading_style.font.size = Pt(12)

# Kod Bloğu Stili
code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = 'Consolas'
code_style.font.size = Pt(8.5)
code_style.font.color.rgb = RGBColor(0xD4, 0xD4, 0xD4)
code_style.paragraph_format.space_before = Pt(6)
code_style.paragraph_format.space_after = Pt(6)
code_style.paragraph_format.line_spacing = 1.15
code_style.paragraph_format.left_indent = Cm(0.5)


def add_code_block(code_text, caption=""):
    if caption:
        cap_para = doc.add_paragraph()
        cap_run = cap_para.add_run(caption)
        cap_run.bold = True
        cap_run.font.size = Pt(10)
        cap_run.font.color.rgb = RGBColor(0x0d, 0x47, 0xa1)
        cap_para.paragraph_format.space_after = Pt(2)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shading = parse_xml('<w:shd %s w:fill="1E1E1E"/>' % nsdecls("w"))
    cell._tc.get_or_add_tcPr().append(shading)
    cell.width = Inches(6.5)
    cell.paragraphs[0].clear()

    lines = code_text.strip().split('\n')
    for i, line in enumerate(lines):
        if i == 0:
            para = cell.paragraphs[0]
        else:
            para = cell.add_paragraph()
        para.style = doc.styles['CodeBlock']
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.1
        run = para.add_run(line if line else ' ')
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0xD4, 0xD4, 0xD4)
    doc.add_paragraph()


def add_bullet(text, bold_prefix=""):
    para = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_bold = para.add_run(bold_prefix)
        run_bold.bold = True
        para.add_run(text)
    else:
        para.add_run(text)
    return para


def add_key_value_table(data, col1_header="Bileşen", col2_header="Teknoloji"):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Shading Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = col1_header
    hdr[1].text = col2_header
    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    for key, value in data:
        row = table.add_row().cells
        row[0].text = key
        row[1].text = value
        for cell in row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


# ====== KAPAK SAYFASI ======
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('ONLİNE HASTANE\nYÖNETİM SİSTEMİ')
run.bold = True
run.font.size = Pt(32)
run.font.color.rgb = RGBColor(0x0d, 0x47, 0xa1)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Yapay Zekâ Destekli Chatbot Entegrasyonu ile')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x42, 0x42, 0x42)
run.italic = True

doc.add_paragraph()

line_para = doc.add_paragraph()
line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
line_run = line_para.add_run('_' * 60)
line_run.font.color.rgb = RGBColor(0x0d, 0x6e, 0xfd)
line_run.font.size = Pt(12)

doc.add_paragraph()

details = [
    ('Proje Türü:', 'Full-Stack Web Uygulaması'),
    ('Ders:', 'Yazılım Mühendisliği'),
    ('Tarih:', 'Nisan 2026'),
    ('Teknolojiler:', 'React.js, Node.js, Express.js, MySQL, OpenAI API'),
]
for label, value in details:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label + ' ')
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    r2 = p.add_run(value)
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_page_break()

# ====== İÇİNDEKİLER ======
doc.add_heading('İçindekiler', level=1)

toc_items = [
    ('1.', 'Giriş ve Proje Genel Bakışı'),
    ('2.', 'Sistem Mimarisi'),
    ('3.', 'Teknoloji Yığını'),
    ('4.', 'Veritabanı Tasarımı'),
    ('5.', 'Arka Uç (Backend) Gerçekleştirimi'),
    ('  5.1', 'Sunucu Giriş Noktası (index.js)'),
    ('  5.2', 'Veritabanı Yapılandırması'),
    ('  5.3', 'JWT Kimlik Doğrulama Ara Yazılımı'),
    ('  5.4', 'Kimlik Doğrulama Denetleyicisi'),
    ('  5.5', 'Randevu Denetleyicisi'),
    ('  5.6', 'E-posta Bildirim Servisi'),
    ('6.', 'Ön Uç (Frontend) Gerçekleştirimi'),
    ('  6.1', 'Uygulama Yönlendirmesi (App.js)'),
    ('  6.2', 'API Servis Katmanı'),
    ('7.', 'Yapay Zekâ Chatbot Entegrasyonu'),
    ('  7.1', 'Chatbot Denetleyicisi - Arka Uç'),
    ('  7.2', 'ChatAssistant - Ön Uç Arayüzü'),
    ('  7.3', 'Araç/Fonksiyon Çağırma Mimarisi'),
    ('8.', 'Güvenlik Özellikleri'),
    ('9.', 'API Uç Noktaları Referansı'),
    ('10.', 'Sonuç'),
]

for num, item in toc_items:
    p = doc.add_paragraph()
    r1 = p.add_run(num + '  ')
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(item)
    r2.font.size = Pt(11)
    if num.startswith('  '):
        p.paragraph_format.left_indent = Cm(1.2)

doc.add_page_break()

# ====== 1. GİRİŞ ======
doc.add_heading('1. Giriş ve Proje Genel Bakışı', level=1)

doc.add_paragraph(
    'Online Hastane Yönetim Sistemi (CarePlus), modern bir hastanenin temel operasyonlarını '
    'dijitalleştirmek ve kolaylaştırmak amacıyla tasarlanmış kapsamlı, full-stack bir web uygulamasıdır. '
    'Sistem, dört farklı kullanıcı türü için rol tabanlı kontrol panelleri sunar: '
    'Yöneticiler, Doktorlar, Hastalar ve Resepsiyonistler - her biri özel işlevsellik ve arayüze sahiptir.'
)

doc.add_paragraph(
    'Bu projenin temel farkı, sanal resepsiyonist olarak görev yapan yapay zekâ destekli bir chatbot '
    'entegrasyonudur. OpenRouter API aracılığıyla 120 Milyar Parametreli Açık Kaynak Büyük Dil Modeli (LLM) '
    'ile güçlendirilmiş olan chatbot, güvenli bir fonksiyon çağırma (tool-use) mimarisi aracılığıyla '
    'hastanenin MySQL veritabanıyla doğrudan etkileşime girerek randevuları otonom olarak '
    'oluşturabilir, yeniden planlayabilir, görüntüleyebilir ve iptal edebilir.'
)

doc.add_heading('Temel Özellikler', level=2)
features = [
    ('Çoklu Rol Sistemi: ', 'Hasta, Doktor, Resepsiyonist ve Yönetici kontrol panelleri ile rol tabanlı erişim kontrolü (RBAC).'),
    ('Randevu Yönetimi: ', 'Çakışma algılama, durum takibi ve e-posta bildirimleri ile tam CRUD işlemleri.'),
    ('Yapay Zekâ Chatbot: ', 'Fonksiyon çağırma özellikli 120 Milyar parametreli LLM kullanarak doğal dil ile randevu yönetimi.'),
    ('E-posta Bildirimleri: ', 'Randevu onayları, iptaller, yeniden planlama ve personel davetleri için otomatik e-posta bildirimleri.'),
    ('Personel Davet Sistemi: ', 'Yönetici, güvenli e-posta tabanlı kayıt bağlantıları aracılığıyla doktor ve resepsiyonistleri davet edebilir.'),
    ('Gerçek Zamanlı Analitik: ', 'Randevu istatistikleri, doktor bazlı dağılım ve günlük trend grafikleri içeren yönetici kontrol paneli.'),
    ('Uluslararasılaştırma: ', 'i18next framework kullanılarak çoklu dil desteği.'),
    ('Güvenlik: ', 'JWT kimlik doğrulama, bcrypt şifre karması, rol tabanlı rota koruması ve SQL enjeksiyon önleme.'),
]
for prefix, text in features:
    add_bullet(text, prefix)

doc.add_page_break()

# ====== 2. SİSTEM MİMARİSİ ======
doc.add_heading('2. Sistem Mimarisi', level=1)

doc.add_paragraph(
    'Uygulama, sunum katmanı (React), iş mantığı katmanı (Express/Node.js) ve veri katmanı (MySQL) '
    'arasında net bir sorumluluk ayrımı bulunan klasik üç katmanlı istemci-sunucu mimarisini takip eder. '
    'Yapay zekâ chatbot, kullanıcı arayüzünü arka uç veritabanı işlemleriyle otonom olarak '
    'köprüleyen dördüncü akıllı bir katman ekler.'
)

doc.add_heading('Mimari Genel Bakışı', level=2)

arch_data = [
    ('Sunum Katmanı', 'React.js 18 SPA, Bootstrap Arayüzü - Port 3000'),
    ('API Katmanı', 'Express.js REST API, JWT kimlik doğrulama ara yazılımı - Port 5000'),
    ('İş Mantığı', '9 Denetleyici: Auth, Appointments, Chatbot, Doctors, Clinics, Users, Messages, Notifications, Medical Records'),
    ('Yapay Zekâ Katmanı', 'OpenRouter LLM API (120B Açık Kaynak model) ile 7 Fonksiyon Aracı'),
    ('E-posta Servisi', 'Nodemailer ile Gmail SMTP - 7 e-posta şablon türü'),
    ('Veri Katmanı', 'MySQL, Bağlantı Havuzu (10 bağlantı) - Parametreli Sorgular'),
]
add_key_value_table(arch_data, 'Katman', 'Teknoloji ve Detaylar')

doc.add_heading('Proje Dosya Yapısı', level=2)
add_code_block('''online-hospital-system/
|-- client/                          # React Ön Uç
|   |-- src/
|   |   |-- components/
|   |   |   |-- admin/               # Yönetici paneli bileşenleri
|   |   |   |-- auth/                # Giriş & Kayıt formları
|   |   |   |-- chat/                # Yapay Zekâ ChatAssistant bileşeni
|   |   |   |-- common/              # Paylaşılan bileşenler
|   |   |   |-- doctor/              # Doktor paneli bileşenleri
|   |   |   |-- layout/              # Başlık, Alt Bilgi, DashboardLayout
|   |   |   |-- patient/             # Hasta paneli bileşenleri
|   |   |   |-- reception/           # Resepsiyonist bileşenleri
|   |   |-- pages/                   # Rota düzeyinde sayfa bileşenleri
|   |   |-- services/
|   |   |   |-- api.js               # Axios HTTP istemcisi & alan servisleri
|   |   |   |-- aiService.js         # Chatbot API servisi
|   |   |-- utils/                   # RBAC araçları (roleGuard.js)
|   |   |-- App.js                   # Yönlendirmeli ana uygulama
|   |   |-- i18n.js                  # Uluslararasılaştırma yapılandırması
|
|-- server/                          # Node.js Arka Uç
|   |-- config/
|   |   |-- db.js                    # MySQL bağlantı havuzu
|   |   |-- email.js                 # Nodemailer e-posta servisi (7 şablon)
|   |-- controllers/
|   |   |-- authController.js        # Kayıt, giriş, davet sistemi
|   |   |-- appointmentsController.js # CRUD + analitik + çakışma algılama
|   |   |-- chatbotController.js     # Yapay Zekâ LLM entegrasyonu, 7 araç
|   |   |-- clinicsController.js     # Bölüm yönetimi
|   |   |-- doctorsController.js     # Doktor sorguları
|   |   |-- usersController.js       # Kullanıcı yönetimi (yönetici)
|   |   |-- messagesController.js    # Mesajlaşma sistemi
|   |   |-- notificationsController.js # Bildirim yönetimi
|   |   |-- medicalRecordsController.js # Tıbbi kayıtlar
|   |-- middleware/
|   |   |-- auth.js                  # JWT doğrulama + rol tabanlı yetkilendirme
|   |-- routes/                      # 10 Express rota dosyası
|   |-- seeder.js                    # Veritabanı tohumlayıcı (klinikler + 30 doktor)
|   |-- index.js                     # Express sunucu giriş noktası
|   |-- .env                         # Ortam değişkenleri''', "Şekil 2.1 - Tam Proje Dosya Yapısı")

doc.add_page_break()

# ====== 3. TEKNOLOJİ YIĞINI ======
doc.add_heading('3. Teknoloji Yığını', level=1)

doc.add_heading('Ön Uç Teknolojileri', level=2)
add_key_value_table([
    ('Arayüz Çerçevesi', 'React.js 18.2'),
    ('Yönlendirme', 'React Router DOM 6.14'),
    ('HTTP İstemcisi', 'Axios 1.4'),
    ('Arayüz Bileşenleri', 'React Bootstrap 2.8 + Bootstrap 5.3'),
    ('Grafikler', 'Chart.js 4.3 + react-chartjs-2'),
    ('İkonlar', 'React Icons 4.12'),
    ('Uluslararasılaştırma', 'i18next + react-i18next'),
    ('Tarih İşleme', 'Moment.js 2.29'),
    ('Kimlik Token Çözümleme', 'jwt-decode 3.1'),
])

doc.add_heading('Arka Uç Teknolojileri', level=2)
add_key_value_table([
    ('Çalışma Zamanı', 'Node.js'),
    ('Web Çerçevesi', 'Express.js 5.1'),
    ('Veritabanı Sürücüsü', 'mysql2 3.14 (Promise tabanlı bağlantı havuzu)'),
    ('Kimlik Doğrulama', 'JSON Web Tokens (jsonwebtoken 9.0)'),
    ('Şifre Karması', 'bcrypt 6.0'),
    ('E-posta Servisi', 'Nodemailer 7.0 (Gmail SMTP)'),
    ('Yapay Zekâ/LLM Entegrasyonu', 'OpenAI SDK 6.34 -> OpenRouter API'),
    ('LLM Modeli', 'openai/gpt-oss-120b:free (120 Milyar parametreli açık kaynak model)'),
    ('Ortam Yapılandırması', 'dotenv 16.5'),
    ('Geliştirme Sunucusu', 'Nodemon 3.1'),
])

doc.add_heading('Veritabanı', level=2)
add_key_value_table([
    ('İlişkisel VT Yönetim Sistemi', 'MySQL'),
    ('Bağlantı Stratejisi', 'Bağlantı Havuzu (10 bağlantı)'),
    ('ORM', 'Parametreli sorgularla ham SQL (SQL enjeksiyonuna karşı güvenli)'),
], 'Bileşen', 'Detaylar')

doc.add_page_break()

# ====== 4. VERİTABANI TASARIMI ======
doc.add_heading('4. Veritabanı Tasarımı', level=1)

doc.add_paragraph(
    'Sistem, aşağıdaki temel tablolara sahip ilişkisel bir MySQL veritabanı kullanır. '
    'Tüm tablolar, referans bütünlüğünü sağlamak için yabancı anahtar ilişkileri aracılığıyla bağlıdır.'
)

doc.add_heading('Varlık İlişki Genel Bakışı', level=2)
add_key_value_table([
    ('users', 'id, username, email, password, role, status, invitation_token, invitation_expires'),
    ('patients', 'id, user_id (FK->users), first_name, last_name, phone'),
    ('doctors', 'id, user_id (FK->users), first_name, last_name, specialization, qualification, phone, clinic_id (FK->clinics)'),
    ('clinics', 'id, name, description, status'),
    ('appointments', 'id, patient_id (FK->patients), doctor_id (FK->doctors), clinic_id (FK->clinics), appointment_date, appointment_time, reason, notes, status'),
    ('notifications', 'id, user_id (FK->users), message, type, related_id'),
    ('receptionists', 'id, user_id (FK->users), first_name, last_name'),
], 'Tablo', 'Temel Sütunlar')

doc.add_heading('Randevu Durum Akışı', level=2)
doc.add_paragraph('Randevular, aşağıdaki geçerli durumlarla iyi tanımlanmış bir yaşam döngüsünü takip eder:')
statuses = [
    ('pending (beklemede): ', 'Hasta form veya chatbot aracılığıyla randevu aldığında başlangıç durumu.'),
    ('pending_assignment (atama bekliyor): ', 'Belirli bir doktor belirtilmeden gönderilmiş; resepsiyonist ataması bekleniyor.'),
    ('confirmed (onaylandı): ', 'Doktor veya resepsiyonist randevuyu onayladı.'),
    ('scheduled (planlandı): ', 'Randevu resmi olarak takvime eklendi.'),
    ('in-progress (devam ediyor): ', 'Hasta şu anda muayene ediliyor.'),
    ('completed (tamamlandı): ', 'Randevu tamamlandı.'),
    ('cancelled (iptal edildi): ', 'Randevu hasta, doktor veya yönetici tarafından iptal edildi.'),
    ('no-show (gelmedi): ', 'Hasta randevuya katılmadı.'),
]
for prefix, desc in statuses:
    add_bullet(desc, prefix)

doc.add_page_break()

# ====== 5. ARKA UÇ GERÇEKLEŞTİRİMİ ======
doc.add_heading('5. Arka Uç (Backend) Gerçekleştirimi', level=1)

doc.add_paragraph(
    'Arka uç, Node.js üzerinde çalışan Express.js 5.1 ile oluşturulmuştur. İş mantığını yöneten '
    'denetleyiciler, API uç noktalarını tanımlayan rotalar, kimlik doğrulama/yetkilendirmeyi '
    'yöneten ara yazılımlar ve veritabanı ile e-posta servisleri için yapılandırma modülleri ile '
    'MVC benzeri bir mimariyi takip eder.'
)

# 5.1 Sunucu Giriş Noktası
doc.add_heading('5.1 Sunucu Giriş Noktası (index.js)', level=2)
doc.add_paragraph(
    'Ana sunucu dosyası Express\'i başlatır, ortam değişkenlerini yükler, ara yazılımları '
    '(CORS, JSON body parser) kaydeder ve tüm rota modüllerini ilgili API yollarına bağlar.'
)

add_code_block('''const express = require('express');
const cors = require('cors');
const dotenv = require('dotenv');
const path = require('path');

// Rota içe aktarmaları
const authRoutes = require('./routes/auth');
const appointmentRoutes = require('./routes/appointments');
const usersRoutes = require('./routes/users');
const notificationsRoutes = require('./routes/notifications');
const messagesRoutes = require('./routes/messages');
const doctorsRoutes = require('./routes/doctors');
const clinicsRoutes = require('./routes/clinics');
const medicalRecordsRoutes = require('./routes/medical-records');
const patientsRoutes = require('./routes/patients');
const chatbotRoutes = require('./routes/chatbot');

dotenv.config({ path: path.resolve(__dirname, '.env') });
const app = express();
const PORT = process.env.PORT || 5000;

// Ara Yazılımlar
app.use(cors());
app.use(express.json());
app.use(express.urlencoded({ extended: true }));

// API Rotaları
app.use('/api/auth', authRoutes);
app.use('/api/appointments', appointmentRoutes);
app.use('/api/users', usersRoutes);
app.use('/api/notifications', notificationsRoutes);
app.use('/api/messages', messagesRoutes);
app.use('/api/doctors', doctorsRoutes);
app.use('/api/clinics', clinicsRoutes);
app.use('/api/medical-records', medicalRecordsRoutes);
app.use('/api/patients', patientsRoutes);
app.use('/api/chatbot', chatbotRoutes);

app.listen(PORT, () => {
  console.log(`Sunucu ${PORT} portunda çalışıyor`);
});''', "Şekil 5.1 - server/index.js: Express Sunucu Başlatma ve Rota Kaydı")

doc.add_paragraph(
    'Sunucu, her biri ayrı bir denetleyiciye eşlenmiş 10 RESTful API rota grubu sunar. '
    'React ön ucunun (port 3000) API ile (port 5000) iletişim kurabilmesi için CORS etkinleştirilmiştir.'
)

# 5.2 Veritabanı Yapılandırması
doc.add_heading('5.2 Veritabanı Yapılandırması (config/db.js)', level=2)
doc.add_paragraph(
    'Veritabanı modülü, Promise desteği ile mysql2 kullanarak bir MySQL bağlantı havuzu oluşturur. '
    'Bu, verimli bağlantı yeniden kullanımı sağlar ve yük altında bağlantı tükenmesini önler.'
)

add_code_block('''const mysql = require('mysql2');
const dotenv = require('dotenv');
const path = require('path');

dotenv.config({ path: path.join(__dirname, '../.env') });

const pool = mysql.createPool({
  host: process.env.DB_HOST || 'localhost',
  port: parseInt(process.env.DB_PORT) || 3306,
  user: process.env.DB_USER || 'root',
  password: process.env.DB_PASSWORD,
  database: process.env.DB_NAME || 'hospital_management',
  waitForConnections: true,
  connectionLimit: 10,
  queueLimit: 0
});

const promisePool = pool.promise();
module.exports = promisePool;''', "Şekil 5.2 - server/config/db.js: Promise Destekli MySQL Bağlantı Havuzu")

doc.add_page_break()

# 5.3 Kimlik Doğrulama Ara Yazılımı
doc.add_heading('5.3 JWT Kimlik Doğrulama Ara Yazılımı', level=2)
doc.add_paragraph(
    'Tüm korumalı rotalar iki katmanlı bir güvenlik ara yazılımından geçer: '
    'authenticateToken, Authorization başlığındaki JWT tokenını doğrular; '
    'authorize ise rol tabanlı erişim kontrolünü uygular.'
)

add_code_block('''const jwt = require('jsonwebtoken');

// JWT tokenını doğrulamak için ara yazılım
const authenticateToken = (req, res, next) => {
  const authHeader = req.headers['authorization'];
  const token = authHeader && authHeader.split(' ')[1];
  
  if (!token) {
    return res.status(401).json({ message: 'Erişim reddedildi. Token sağlanmadı.' });
  }

  try {
    const decoded = jwt.verify(token, process.env.JWT_SECRET);
    req.user = decoded;
    next();
  } catch (error) {
    return res.status(403).json({ message: 'Geçersiz token.' });
  }
};

// Rol tabanlı yetkilendirme ara yazılımı
const authorize = (roles = []) => {
  if (typeof roles === 'string') {
    roles = [roles];
  }

  return (req, res, next) => {
    if (roles.length === 0) return next();
    if (!req.user || !roles.includes(req.user.role)) {
      return res.status(403).json({ 
        message: 'Yasak: Bu işlem için yetkiniz yok' 
      });
    }
    next();
  };
};

module.exports = { authenticateToken, authorize };''', "Şekil 5.3 - server/middleware/auth.js: JWT Token Doğrulama ve Rol Yetkilendirmesi")

# 5.4 Kimlik Doğrulama Denetleyicisi
doc.add_heading('5.4 Kimlik Doğrulama Denetleyicisi (Öne Çıkanlar)', level=2)
doc.add_paragraph(
    'Kimlik doğrulama denetleyicisi; kullanıcı kaydı, giriş, e-posta doğrulama, şifre yönetimi '
    've personel davet sistemini yönetir. Aşağıda temel fonksiyonlar verilmiştir:'
)

doc.add_heading('E-posta Doğrulamalı Kullanıcı Kaydı', level=3)
add_code_block('''exports.register = async (req, res) => {
  try {
    const { username, email, password, role, first_name, last_name, 
            phone, specialization, qualification, verificationCode } = req.body;

    // E-posta doğrulama kontrolü
    if (process.env.EMAIL_VERIFICATION_ENABLED === 'true') {
      if (!verificationCodes[email] || !verificationCodes[email].verified) {
        // E-posta ile doğrulama kodu gönder
      }
    }

    // Rol kısıtlaması: Yalnızca yöneticiler hasta dışı rolleri kaydedebilir
    if (role !== 'patient') {
      const authHeader = req.headers['authorization'];
      const token = authHeader && authHeader.split(' ')[1];
      const decoded = jwt.verify(token, process.env.JWT_SECRET);
      if (decoded.role !== 'admin') 
        return res.status(403).json({ message: 'Yönetici yetkileri gerekli.' });
    }

    // bcrypt ile şifre karması (10 tuz turu)
    const salt = await bcrypt.genSalt(10);
    const hashedPassword = await bcrypt.hash(password, salt);

    // İşlemsel ekleme: kullanıcı + role özel profil
    await db.query('START TRANSACTION');
    const [result] = await db.query(
      'INSERT INTO users (username, email, password, role) VALUES (?, ?, ?, ?)',
      [username, email, hashedPassword, role]
    );

    if (role === 'doctor') {
      await db.query(
        'INSERT INTO doctors (user_id, first_name, last_name, specialization) VALUES (...)',
        [result.insertId, first_name, last_name, specialization]
      );
    } else if (role === 'patient') {
      await db.query(
        'INSERT INTO patients (user_id, first_name, last_name, phone) VALUES (...)',
        [result.insertId, first_name, last_name, phone]
      );
    }

    const token = jwt.sign(
      { id: result.insertId, username, email, role },
      process.env.JWT_SECRET, { expiresIn: '1d' }
    );

    await db.query('COMMIT');
    res.status(201).json({ message: 'Kullanıcı başarıyla kaydedildi', token });
  } catch (error) {
    await db.query('ROLLBACK'); throw error;
  }
};''', "Şekil 5.4a - authController.js: İşlem ve Rol Yönetimli Kullanıcı Kaydı")

doc.add_page_break()

doc.add_heading('JWT Oluşturmalı Kullanıcı Girişi', level=3)
add_code_block('''exports.login = async (req, res) => {
  try {
    const { username, password } = req.body;

    const [users] = await db.query(
      'SELECT * FROM users WHERE username = ? OR email = ?',
      [username, username]  // Kullanıcı adı veya e-posta ile giriş izni
    );

    if (users.length === 0) 
      return res.status(401).json({ message: 'Geçersiz kimlik bilgileri' });

    const user = users[0];
    const isPasswordValid = await bcrypt.compare(password, user.password);
    if (!isPasswordValid) 
      return res.status(401).json({ message: 'Geçersiz kimlik bilgileri' });

    // Kullanıcı verileriyle JWT tokenı oluştur (24 saat süre)
    const token = jwt.sign(
      { id: user.id, username: user.username, email: user.email, role: user.role },
      process.env.JWT_SECRET || 'hospital_system_jwt_secret_key',
      { expiresIn: '1d' }
    );

    res.json({ 
      token, 
      user: { id: user.id, username: user.username, role: user.role } 
    });
  } catch (error) {
    res.status(500).json({ message: 'Sunucu hatası' });
  }
};''', "Şekil 5.4b - authController.js: bcrypt Şifre Karşılaştırma ve JWT Token ile Giriş")

# 5.5 Randevu Denetleyicisi
doc.add_heading('5.5 Randevu Denetleyicisi (Öne Çıkanlar)', level=2)
doc.add_paragraph(
    'Randevu denetleyicisi, en büyük arka uç modülüdür (~570 satır). Atomik çakışma kontrolü ile oluşturma, '
    'kademeli e-posta bildirimleri ile durum güncellemeleri, analitik ve sahiplik doğrulamasını yönetir.'
)

add_code_block('''exports.createAppointment = async (req, res) => {
  try {
    const { doctorId, clinicId, date, time, reason, notes, patientUserId } = req.body;
    const isReceptionist = req.user.role === 'receptionist';
    const targetUserId = isReceptionist && patientUserId ? patientUserId : req.user.id;

    // Hasta Profilini Çözümle
    const [patients] = await db.query(
      'SELECT id, first_name, last_name FROM patients WHERE user_id = ?', 
      [targetUserId]
    );
    if (patients.length === 0) 
      return res.status(404).json({ message: 'Hasta kaydı bulunamadı.' });

    // Atomik Çakışma Kontrolü & Ekleme (İşlem)
    const conn = await db.getConnection();
    await conn.beginTransaction();

    if (doctorId) {
      const [conflict] = await conn.query(
        `SELECT id FROM appointments
         WHERE doctor_id = ? AND appointment_date = ?
           AND TIME_FORMAT(appointment_time,'%H:%i') = TIME_FORMAT(?,'%H:%i')
           AND status NOT IN ('cancelled','no-show')`,
        [doctorId, date, formattedTime]
      );
      if (conflict.length > 0) {
        await conn.rollback(); conn.release();
        return res.status(409).json({ message: 'Zaman dilimi zaten ayrılmış.' });
      }
    }

    const [result] = await conn.query(
      'INSERT INTO appointments (...) VALUES (?, ?, ?, ?, ?, ?, ?, ?)',
      [patient.id, doctorId, clinicId, date, formattedTime, reason, notes, 'pending']
    );

    await conn.commit(); conn.release();

    // Doktora engelsiz e-posta bildirimi
    emailService.sendDoctorNewAppointmentEmail(...).catch(console.error);

    return res.status(201).json({ 
      message: 'Randevu başarıyla oluşturuldu', 
      id: result.insertId 
    });
  } catch (error) { ... }
};''', "Şekil 5.5 - appointmentsController.js: Çakışma Algılamalı Atomik Randevu Oluşturma")

doc.add_page_break()

# 5.6 E-posta Servisi
doc.add_heading('5.6 E-posta Bildirim Servisi', level=2)
doc.add_paragraph(
    'Sistem, otomatik e-posta bildirimleri göndermek için Gmail SMTP ile Nodemailer kullanır. '
    'E-posta servisi, her biri profesyonel HTML şablonları içeren 7 farklı e-posta türü sağlar:'
)

email_types = [
    'E-posta Doğrulama (kayıt için 6 haneli OTP)',
    'Personel Daveti (güvenli kurulum bağlantısı ile, 48 saat süreli)',
    'Randevu Beklemede (hastaya onay)',
    'Randevu Onayı (doktor, tarih, saat, klinik detayları ile)',
    'Randevu Yeniden Planlandı (güncellenmiş detay bildirimi)',
    'Randevu İptali (iptal nedeni ile)',
    'Doktora Yeni Randevu (atanan doktora bildirim)',
]
for t in email_types:
    add_bullet(t)

add_code_block('''const nodemailer = require('nodemailer');

const transporter = nodemailer.createTransport({
  service: 'gmail',
  auth: {
    user: process.env.SMTP_USER,
    pass: process.env.SMTP_PASS  // Gmail Uygulama Şifresi
  }
});

const sendAppointmentConfirmationEmail = async (
  patientEmail, patientName, doctorName, date, time, clinic
) => {
  const mailOptions = {
    from: SMTP_USER,
    to: patientEmail,
    subject: 'Randevu Onaylandı - Hastane Yönetim Sistemi',
    html: `
      <div style="font-family: Arial; max-width: 600px; margin: auto;">
        <h2>Randevu Onayı</h2>
        <p>Sayın ${patientName},</p>
        <p>Randevunuz onaylanmıştır:</p>
        <ul>
          <li><strong>Doktor:</strong> ${doctorName}</li>
          <li><strong>Tarih:</strong> ${date}</li>
          <li><strong>Saat:</strong> ${time}</li>
          <li><strong>Klinik:</strong> ${clinic}</li>
        </ul>
        <p>Lütfen randevunuzdan 15 dakika önce gelin.</p>
      </div>`
  };
  await transporter.sendMail(mailOptions);
};''', "Şekil 5.6 - config/email.js: Nodemailer SMTP Kurulumu ve E-posta Şablonu Örneği")

doc.add_page_break()

# ====== 6. ÖN UÇ GERÇEKLEŞTİRİMİ ======
doc.add_heading('6. Ön Uç (Frontend) Gerçekleştirimi', level=1)

doc.add_paragraph(
    'Ön uç, istemci taraflı yönlendirme için React Router, arayüz bileşenleri için React Bootstrap '
    've API iletişimi için Axios kullanan bir React.js 18 tek sayfa uygulamasıdır (SPA). '
    'Uygulama, paylaşılan düzenler ve role özel sayfalar içeren bileşen tabanlı bir mimari kullanır.'
)

# 6.1 App.js
doc.add_heading('6.1 Uygulama Yönlendirmesi ve Rol Tabanlı Erişim (App.js)', level=2)
doc.add_paragraph(
    'Ana App bileşeni, eksiksiz bir rol tabanlı yönlendirme sistemi uygular. ProtectedRoute sarmalayıcı '
    'bileşeni hem kimlik doğrulamayı (/login\'e yönlendirme) hem de yetkilendirmeyi (yetkisiz roller için '
    '403 ekranı) zorlar. Her rolün kendi rota grubu ve ayrılmış kontrol paneli düzeni vardır.'
)

add_code_block('''function App() {
  const [isAuthenticated, setIsAuthenticated] = useState(false);
  const [user, setUser] = useState(null);

  // ProtectedRoute - kimlik doğrulama + rol yetkilendirmesini zorlar
  const ProtectedRoute = ({ children, allowedRoles = [] }) => {
    if (!isAuthenticated) return <Navigate to="/login" replace />;
    if (allowedRoles.length > 0 && !allowedRoles.includes(user?.role)) 
      return <Unauthorized />;
    return children;
  };

  return (
    <Router>
      <DashboardLayout user={user} logout={logout}>
        <Routes>
          {/* Yönetici Rotaları */}
          <Route path="/admin" element={
            <ProtectedRoute allowedRoles={['admin']}>
              <AdminDashboard />
            </ProtectedRoute>
          } />
          <Route path="/admin/doctors" element={...} />
          <Route path="/admin/appointments" element={...} />
          <Route path="/admin/users" element={...} />
          <Route path="/admin/departments" element={...} />

          {/* Doktor Rotaları */}
          <Route path="/doctor" element={
            <ProtectedRoute allowedRoles={['doctor']}>
              <DoctorDashboard />
            </ProtectedRoute>
          } />
          <Route path="/doctor/appointments" element={...} />

          {/* Hasta Rotaları */}
          <Route path="/patient" element={...} />
          <Route path="/patient/appointments/book" element={...} />

          {/* Resepsiyon Rotaları */}
          <Route path="/reception" element={...} />
          <Route path="/reception/appointments" element={...} />

          {/* Akıllı Ana Sayfa Yönlendirmesi */}
          <Route path="/" element={
            <Navigate to={getHomePath(user?.role)} replace />
          } />
        </Routes>
      </DashboardLayout>
    </Router>
  );
}''', "Şekil 6.1 - client/src/App.js: Rol Tabanlı Korumalı Yönlendirme Sistemi")

doc.add_page_break()

# 6.2 API Servisi
doc.add_heading('6.2 Merkezi API Servis Katmanı', level=2)
doc.add_paragraph(
    'Tüm HTTP iletişimi, Axios tabanlı bir API servis modülü aracılığıyla merkezileştirilmiştir. '
    'İstek interceptor\'ları otomatik olarak JWT tokenlarını ekler ve yanıt interceptor\'ları '
    'oturum süresi dolması (401 -> otomatik giriş sayfasına yönlendirme) gibi genel hata durumlarını yönetir.'
)

add_code_block('''import axios from 'axios';

const API_BASE_URL = process.env.REACT_APP_API_URL || 'http://localhost:5000/api';

const api = axios.create({
  baseURL: API_BASE_URL,
  headers: { 'Content-Type': 'application/json' }
});

// Her giden isteğe otomatik JWT tokenı ekleme
api.interceptors.request.use((config) => {
  const token = localStorage.getItem('token');
  if (token) config.headers.Authorization = `Bearer ${token}`;
  return config;
});

// Süresi dolmuş oturumları otomatik yönetme (401 -> giriş sayfasına yönlendirme)
api.interceptors.response.use(
  (response) => response,
  (error) => {
    if (error.response?.status === 401) {
      localStorage.removeItem('token');
      localStorage.removeItem('user');
      window.location.href = '/login';
    }
    return Promise.reject(error);
  }
);

// Alan Servis Nesneleri
export const appointmentService = {
  getAll:         () => api.get('/appointments'),
  create:        (data) => api.post('/appointments', data),
  update:    (id, data) => api.patch(`/appointments/${id}/status`, data),
  getAnalytics:  () => api.get('/appointments/analytics'),
};

export const chatbotService = {
  sendMessage: (message) => api.post('/chatbot/message', { message }),
};''', "Şekil 6.2 - client/src/services/api.js: Axios Interceptor'lar ve Alan Servisleri")

doc.add_page_break()

# ====== 7. YAPAY ZEKÂ CHATBOT ENTEGRASYONU ======
doc.add_heading('7. Yapay Zekâ Chatbot Entegrasyonu', level=1)

doc.add_paragraph(
    'Bu sistemin en yenilikçi özelliği, sanal hastane resepsiyonisti olarak görev yapan yapay zekâ '
    'destekli chatbot\'tur. Geleneksel kural tabanlı chatbot\'ların aksine, bu uygulama OpenRouter API '
    'aracılığıyla 120 Milyar Parametreli Açık Kaynak Büyük Dil Modeli (LLM) kullanarak doğal dil '
    'anlayışı ve otonom görev yürütme sağlar.'
)

doc.add_heading('Nasıl Çalışır', level=2)
doc.add_paragraph(
    'Chatbot mimarisi, LLM\'nin sadece metin üretmediği - sunucunun veritabanı işlemlerini '
    'yürütmesini aktif olarak isteyebildiği bir "Araç Kullanımı / Fonksiyon Çağırma" modeline dayanır. '
    'Akış şu şekilde çalışır:'
)

steps = [
    'Hasta doğal dil mesajı gönderir (örn. "Yarın bir kardiyolog görmek istiyorum").',
    'Mesaj, konuşma geçmişi ile birlikte arka uç chatbot denetleyicisine gönderilir.',
    'Denetleyici, konuşmayı tanımlı araçlarla birlikte OpenRouter LLM API\'sine iletir.',
    'LLM isteği analiz eder ve yapılandırılmış bir JSON araç çağrısı çıktısı verir (örn. get_clinics).',
    'Sunucu veritabanı sorgusunu yürütür ve sonucu LLM\'ye geri gönderir.',
    'LLM birden fazla araç çağrısını zincirleme yapabilir (get_clinics -> get_doctors -> check_availability -> book).',
    'Tüm veriler toplandığında, LLM hastaya son doğal dil yanıtını oluşturur.',
]
for i, step in enumerate(steps, 1):
    add_bullet(step, 'Adım %d: ' % i)

doc.add_page_break()

# 7.1 Chatbot Denetleyicisi
doc.add_heading('7.1 Chatbot Denetleyicisi - Arka Uç (chatbotController.js)', level=2)

doc.add_heading('LLM Başlatma ve Sistem İstemi', level=3)
doc.add_paragraph(
    'Denetleyici, OpenRouter API uç noktasına işaret eden OpenAI SDK\'sını başlatır. '
    'Kapsamlı bir sistem talimatı, yapay zekânın kişiliğini ve operasyonel kurallarını tanımlar.'
)

add_code_block('''const OpenAI = require('openai');
const db = require('../config/db');

// OpenRouter İstemci Başlatma (Özel base URL ile OpenAI SDK)
const openai = new OpenAI({
  baseURL: 'https://openrouter.ai/api/v1',
  apiKey: process.env.OPENROUTER_API_KEY || 'MISSING_KEY'
});

const systemInstruction = `Sen CarePlus Asistanı'sın, bir hastane için yardımcı 
bir yapay zekâ sanal resepsiyonistisin. 
Temel görevin hastaların tıbbi randevularını oluşturma, yeniden planlama, 
görüntüleme ve iptal etmelerine yardımcı olmaktır.
Veritabanıyla etkileşim kurmak için arka uç fonksiyonel araçlarına erişimin var.

RANDEVU OLUŞTURMA İÇİN KRİTİK KURALLAR:
1. Randevu alırken Klinik, Doktor, Tarih ve Saati toplamalısın.
2. Bölümleri göstermek için 'get_clinics', ardından 'get_doctors' çağır.
3. Randevuyu kesinleştirmeden önce MUTLAKA 'check_availability' çağırmalısın.
4. Uygunluğu doğruladıktan sonra kullanıcıdan onay iste.

YENİDEN PLANLAMA İÇİN KRİTİK KURALLAR:
1. Her zaman önce mevcut randevuları görmek için 'get_my_appointments' çağır.
2. Yeni tarih için doktorun 'check_availability' kontrolünü MUTLAKA yapmalısın.
3. Yalnızca zaman dilimi müsaitse 'reschedule_appointment' çağır.`;''', "Şekil 7.1a - chatbotController.js: OpenRouter LLM İstemcisi ve Sistem Talimatı")

doc.add_heading('Araç/Fonksiyon Tanımları', level=3)
doc.add_paragraph(
    'LLM\'ye çağırabileceği 7 aracın yapılandırılmış bir dizisi sağlanır. Her aracın '
    'bir adı, açıklaması ve parametreleri için JSON şeması vardır:'
)

add_code_block('''const tools = [
  {
    type: 'function',
    function: {
      name: 'get_clinics',
      description: 'Tüm hastane kliniklerinin/bölümlerinin listesini al.',
      parameters: { type: 'object', properties: {} } 
    }
  },
  {
    type: 'function',
    function: {
      name: 'get_doctors',
      description: 'Belirli bir klinik için doktorları al.',
      parameters: {
        type: 'object',
        properties: {
          clinicId: { type: 'string', description: 'Klinik ID.' }
        },
        required: ['clinicId']
      }
    }
  },
  {
    type: 'function',
    function: {
      name: 'check_availability',
      description: 'Bir doktor için belirli bir tarihteki müsait zaman dilimlerini ara.',
      parameters: {
        type: 'object',
        properties: {
          doctorId: { type: 'string' },
          date: { type: 'string', description: 'YYYY-MM-DD formatı.' }
        },
        required: ['doctorId', 'date']
      }
    }
  },
  {
    type: 'function',
    function: {
      name: 'book_appointment',
      description: 'Yeni bir randevu oluştur.',
      parameters: {
        type: 'object',
        properties: {
          clinicId: { type: 'string' }, doctorId: { type: 'string' },
          date: { type: 'string' },     time: { type: 'string' }
        },
        required: ['clinicId', 'doctorId', 'date', 'time']
      }
    }
  },
  // ... get_my_appointments, cancel_appointment, reschedule_appointment
];''', "Şekil 7.1b - chatbotController.js: LLM Araç Tanımları (JSON Şeması)")

doc.add_page_break()

doc.add_heading('Çok Adımlı Araç Zincirleme ile Mesaj İşleme Döngüsü', level=3)
doc.add_paragraph(
    'Chatbot\'un çekirdeği, LLM son bir metin yanıtı üretene kadar araç sonuçlarını '
    'LLM\'ye geri göndermeye devam eden bir while döngüsüdür. Bu, yapay zekânın tek bir '
    'kullanıcı isteği içinde birden fazla veritabanı işlemini otonom olarak zincirlemesini sağlar.'
)

add_code_block('''// Sohbet belleği için bellek içi oturum deposu
const activeChats = {};

exports.processMessage = async (req, res) => {
  const userId = req.user.id;
  const { message } = req.body;

  // Yeni kullanıcılar için konuşma geçmişini başlat
  if (!activeChats[userId]) {
    activeChats[userId] = [
      { role: 'system', content: systemInstruction }
    ];
  }
  
  activeChats[userId].push({ role: 'user', content: message });

  let isMakingToolCalls = true;
  let finalContent = "";

  // İşleme döngüsü: birden fazla araç çağrısı turunu otomatik yönet
  while (isMakingToolCalls) {
    const response = await openai.chat.completions.create({
      model: 'openai/gpt-oss-120b:free', 
      messages: activeChats[userId],
      tools: tools,
    });

    const responseMessage = response.choices[0].message;
    activeChats[userId].push(responseMessage);

    if (responseMessage.tool_calls && responseMessage.tool_calls.length > 0) {
      // Her aracı yürüt ve sonuçları LLM'ye geri besle
      for (const call of responseMessage.tool_calls) {
        const args = JSON.parse(call.function.arguments);
        const toolResult = await executeTool(
          call.function.name, args, userId
        );
        
        activeChats[userId].push({
          tool_call_id: call.id,
          role: 'tool',
          name: call.function.name,
          content: JSON.stringify(toolResult)
        });
      }
      // Döngü devam eder - LLM araç sonuçlarını işleyecek
    } else {
      // Artık araç çağrısı yok - son metin yanıtı hazır
      isMakingToolCalls = false;
      finalContent = responseMessage.content;
    }
  }
  
  return res.json({ text: finalContent });
};''', "Şekil 7.1c - chatbotController.js: Bellekli Çok Adımlı Araç Çağırma Döngüsü")

doc.add_page_break()

doc.add_heading('Araç Yürütücü Fonksiyonları', level=3)
doc.add_paragraph(
    'Her araç, belirli bir veritabanı işlemine karşılık gelir. LLM asla doğrudan SQL yazmaz - '
    'yapılandırılmış JSON parametreleri çıktılar ve sunucu parametreli sorguları güvenle yürütür.'
)

add_code_block('''async function executeTool(name, args, userId) {
  switch(name) {
    case 'get_clinics': {
      const [clinics] = await db.query('SELECT id, name FROM clinics');
      return { clinics };
    }
    
    case 'get_doctors': {
      const [doctors] = await db.query(
        'SELECT id, first_name, last_name, specialization ' +
        'FROM doctors WHERE clinic_id = ?',
        [args.clinicId]
      );
      return { doctors };
    }
    
    case 'check_availability': {
      const [booked] = await db.query(
        `SELECT TIME_FORMAT(appointment_time,'%H:%i') as t 
         FROM appointments 
         WHERE doctor_id = ? AND appointment_date = ? 
           AND status NOT IN ('cancelled')`,
        [args.doctorId, args.date]
      );
      const bookedTimes = booked.map(r => r.t);
      // 09:00-17:00 arası 30 dakikalık tüm zaman dilimlerini oluştur
      const allSlots = [];
      for (let h = 9; h < 17; h++) {
        allSlots.push(`${String(h).padStart(2,'0')}:00`);
        allSlots.push(`${String(h).padStart(2,'0')}:30`);
      }
      const available = allSlots.filter(s => !bookedTimes.includes(s));
      return { available_slots: available };
    }
    
    case 'book_appointment': {
      const [patients] = await db.query(
        'SELECT id FROM patients WHERE user_id = ?', [userId]
      );
      // Ekleme öncesi çakışmayı tekrar kontrol et
      const [conflict] = await db.query(
        "SELECT id FROM appointments WHERE doctor_id = ? " +
        "AND appointment_date = ? AND status NOT IN ('cancelled')",
        [args.doctorId, args.date]
      );
      if (conflict.length > 0) 
        return { status: 'error', message: 'Zaman dilimi dolu.' };

      await db.query(
        'INSERT INTO appointments (...) VALUES (?, ?, ?, ?, ?, ?, ?)',
        [patients[0].id, args.doctorId, args.clinicId, 
         args.date, args.time, 'Chatbot AI Randevusu', 'pending']
      );
      return { status: 'success', message: 'Randevu oluşturuldu.' };
    }
    
    case 'cancel_appointment': {
      await db.query(
        "UPDATE appointments SET status = 'cancelled' WHERE id = ?",
        [args.appointmentId]
      );
      return { status: 'success' };
    }
    
    case 'reschedule_appointment': {
      // Yeni zaman diliminin müsait olduğunu doğrula, ardından güncelle
      const [conflict] = await db.query(...);
      if (conflict.length > 0) 
        return { status: 'error', message: 'Zaman dilimi müsait değil.' };
      await db.query(
        "UPDATE appointments SET appointment_date = ?, " +
        "appointment_time = ?, status = 'pending' WHERE id = ?",
        [args.newDate, args.newTime, args.appointmentId]
      );
      return { status: 'success', message: 'Yeniden planlandı.' };
    }
  }
}''', "Şekil 7.1d - chatbotController.js: Araç Yürütücü - Güvenli Veritabanı İşlemleri")

doc.add_page_break()

# 7.2 ChatAssistant Ön Uç
doc.add_heading('7.2 ChatAssistant - Ön Uç Arayüz Bileşeni', level=2)
doc.add_paragraph(
    'ChatAssistant, tüm hasta sayfalarında görünen yüzen bir React bileşenidir. '
    'Mesaj balonu arayüzü, yazıyor göstergesi, markdown oluşturma ve localStorage aracılığıyla '
    'kalıcı mesaj geçmişi özelliklerini sunar.'
)

add_code_block('''import React, { useState, useEffect } from 'react';
import { Card, Button, Form, InputGroup } from 'react-bootstrap';
import { FaRobot, FaTimes, FaCommentMedical, FaPaperPlane } from 'react-icons/fa';
import { processChatMessage } from '../../services/aiService';

const ChatAssistant = () => {
  const [isOpen, setIsOpen] = useState(false);
  const [input, setInput] = useState('');
  const [isTyping, setIsTyping] = useState(false);
  
  // localStorage aracılığıyla kalıcı mesaj geçmişi
  const [messages, setMessages] = useState(() => {
    const saved = localStorage.getItem('chat_messages');
    return saved ? JSON.parse(saved) : [
      { sender: 'bot', 
        text: "Merhaba! Ben CarePlus Yapay Zekâ Asistanınızım." }
    ];
  });

  // Her değişiklikte mesajları localStorage'a senkronize et
  useEffect(() => {
    localStorage.setItem('chat_messages', JSON.stringify(messages));
  }, [messages]);

  const handleSend = async () => {
    const displayMsg = input.trim();
    if (!displayMsg) return;

    // Kullanıcı mesajını anında göster
    setMessages(prev => [...prev, { sender: 'user', text: displayMsg }]);
    setInput('');
    setIsTyping(true);

    try {
      const response = await processChatMessage(null, displayMsg, null);
      setTimeout(() => {
        setIsTyping(false);
        setMessages(prev => [...prev, { 
          sender: 'bot', text: response.text 
        }]);
      }, 500);  // Doğal his için küçük gecikme
    } catch (error) {
      setIsTyping(false);
      setMessages(prev => [...prev, { 
        sender: 'bot', text: 'Hata. Lütfen tekrar deneyin.' 
      }]);
    }
  };

  // Kalın metin ve madde işaretleri için Markdown ayrıştırıcı
  const parseMarkdown = (text) => {
    if (!text) return '';
    let html = text.replace(/\\*\\*(.*?)\\*\\*/g, '<b>$1</b>');
    html = html.replace(/^\\* (.*$)/gim, '  $1');
    return html;
  };

  return (
    <>
      {/* Yüzen tetikleme düğmesi (sağ alt köşe) */}
      {!isOpen && (
        <Button variant="primary" 
                className="rounded-circle shadow-lg position-fixed"
                style={{ width: 60, height: 60, bottom: 30, right: 30 }}
                onClick={() => setIsOpen(true)}>
          <FaCommentMedical size={26} />
        </Button>
      )}

      {/* Sohbet penceresi: başlık + mesajlar + giriş */}
      {isOpen && (
        <Card className="position-fixed shadow" 
              style={{ width: 380, height: 600, 
                       bottom: 30, right: 30, zIndex: 1050 }}>
          {/* Bot ikonu ve başlık ile üst bilgi */}
          {/* Otomatik kaydırmalı mesaj alanı */}
          {/* Yazıyor göstergesi (3 zıplayan nokta) */}
          {/* Gönder düğmeli giriş alanı */}
        </Card>
      )}
    </>
  );
};''', "Şekil 7.2 - ChatAssistant.js: Yüzen Yapay Zekâ Sohbet Arayüzü Bileşeni")

doc.add_page_break()

# 7.3 Araç Mimarisi Özeti
doc.add_heading('7.3 Araç/Fonksiyon Çağırma Mimarisi Özeti', level=2)

add_key_value_table([
    ('get_clinics', 'Clinics tablosundan tüm hastane bölümlerini okur.'),
    ('get_doctors', 'Doktorları clinic_id\'ye göre filtreler; ad ve uzmanlık döndürür.'),
    ('check_availability', '09:00-17:00 arası 30 dakikalık dilimler oluşturur, dolu olanları filtreler.'),
    ('book_appointment', 'Çakışmaları tekrar kontrol eder, ardından yeni bekleyen randevu ekler (INSERT).'),
    ('get_my_appointments', 'Giriş yapmış hasta için randevuları doktorlar ve kliniklerle birleştirir (JOIN).'),
    ('cancel_appointment', 'Randevu durumunu ID\'ye göre "cancelled" olarak ayarlar.'),
    ('reschedule_appointment', 'Yeni zaman dilimini tekrar kontrol eder, ardından tarih/saati günceller (UPDATE).'),
], 'Araç Adı', 'Veritabanı İşlemi')

doc.add_page_break()

# ====== 8. GÜVENLİK ÖZELLİKLERİ ======
doc.add_heading('8. Güvenlik Özellikleri', level=1)

security_items = [
    ('JWT Kimlik Doğrulama: ', 'Tüm API uç noktaları (giriş/kayıt hariç) Authorization başlığında geçerli bir JWT tokenı gerektirir. Tokenlar 24 saat sonra süreleri dolar.'),
    ('bcrypt Şifre Karması: ', 'Kullanıcı şifreleri, depolanmadan önce bcrypt (10 tur) ile tuzlanır ve karma haline getirilir. Düz metin şifreler asla saklanmaz.'),
    ('Rol Tabanlı Erişim Kontrolü: ', 'Hem ön uç rotaları hem de arka uç uç noktaları rol kontrollerini uygular. Yalnızca yöneticiye özel işlemler her iki katmanda da korunur.'),
    ('SQL Enjeksiyon Önleme: ', 'Tüm veritabanı sorguları parametreli yer tutucular (?) kullanır. Yapay zekâ chatbot asla SQL üretmez - yapılandırılmış JSON parametreleri çıktılar.'),
    ('İşlemsel Bütünlük: ', 'Kritik operasyonlar, kısmi yazmayı önlemek için MySQL işlemlerini (START TRANSACTION / COMMIT / ROLLBACK) kullanır.'),
    ('Personel Davet Güvenliği: ', 'Personel kayıt sistemi, 48 saatlik sona erme süresine sahip kriptografik olarak rastgele tokenlar (crypto.randomBytes) kullanır.'),
    ('Hız Sınırlama Yönetimi: ', 'Chatbot denetleyicisi, yapay zekâ sağlayıcısından gelen HTTP 429 yanıtlarını zarfice yönetir ve kullanıcı dostu bir mesaj döndürür.'),
    ('Oturum İzolasyonu: ', 'Her hastanın chatbot konuşması izole bir bellek içi oturumda saklanır (activeChats[userId]).'),
]

for prefix, text in security_items:
    add_bullet(text, prefix)

doc.add_page_break()

# ====== 9. API UÇ NOKTALARI REFERANSI ======
doc.add_heading('9. API Uç Noktaları Referansı', level=1)

doc.add_heading('Kimlik Doğrulama', level=2)
add_key_value_table([
    ('POST /api/auth/register', 'Yeni kullanıcı kaydı (hasta: herkese açık, diğerleri: yalnızca yönetici)'),
    ('POST /api/auth/login', 'Giriş yap ve JWT tokenı al'),
    ('GET /api/auth/me', 'Mevcut kullanıcı profilini al'),
    ('PATCH /api/auth/change-password', 'Şifre güncelle (mevcut şifre gerektirir)'),
    ('POST /api/auth/invite', 'E-posta ile personel davet et (yalnızca yönetici)'),
    ('GET /api/auth/verify-invite', 'Davet tokenını doğrula'),
    ('POST /api/auth/setup-invited', 'Davetli hesap kurulumunu tamamla'),
], 'Uç Nokta', 'Açıklama')

doc.add_heading('Randevular', level=2)
add_key_value_table([
    ('GET /api/appointments', 'Randevuları al (kullanıcı rolüne göre otomatik filtrelenir)'),
    ('POST /api/appointments', 'Yeni randevu oluştur'),
    ('PATCH /api/appointments/:id/status', 'Durum güncelle / yeniden planla / doktor ata'),
    ('DELETE /api/appointments/:id', 'Randevu sil'),
    ('GET /api/appointments/analytics', 'Kontrol paneli analitiğini al (yönetici)'),
    ('GET /api/appointments/availability', 'Zaman dilimi uygunluğunu kontrol et'),
], 'Uç Nokta', 'Açıklama')

doc.add_heading('Doktorlar, Klinikler ve Hastalar', level=2)
add_key_value_table([
    ('GET /api/doctors', 'Tüm doktorları al'),
    ('GET /api/doctors/clinic/:clinicId', 'Bölüme göre doktorları al'),
    ('GET /api/clinics', 'Tüm bölümleri al'),
    ('POST /api/clinics', 'Bölüm oluştur (yönetici)'),
    ('PUT /api/clinics/:id', 'Bölüm güncelle (yönetici)'),
    ('DELETE /api/clinics/:id', 'Bölüm sil (yönetici)'),
    ('GET /api/users', 'Tüm kullanıcıları listele (yönetici)'),
    ('GET /api/users/stats', 'Kontrol paneli istatistikleri (yönetici)'),
], 'Uç Nokta', 'Açıklama')

doc.add_heading('Yapay Zekâ Chatbot', level=2)
add_key_value_table([
    ('POST /api/chatbot/message', 'Yapay zekâ asistanına mesaj gönder (kimliği doğrulanmış hastalar)'),
], 'Uç Nokta', 'Açıklama')

doc.add_page_break()

# ====== 10. SONUÇ ======
doc.add_heading('10. Sonuç', level=1)

doc.add_paragraph(
    'Online Hastane Yönetim Sistemi (CarePlus), sağlık bilgi sistemlerine kapsamlı, üretim düzeyinde '
    'bir yaklaşım sergilemektedir. Modern web teknolojileri - ön uçta React.js, arka uçta Node.js/Express '
    've veri kalıcılığı için MySQL - ile oluşturulan sistem, temiz kod mimarisi ve sağlam güvenlik '
    'uygulamalarını sürdürürken tüm temel hastane yönetimi işlevselliklerini başarıyla uygulamaktadır.'
)

doc.add_paragraph(
    'Yapay zekâ chatbot entegrasyonu, projenin teknik açıdan en iddialı bileşenini temsil eder. '
    'Fonksiyon çağırma (tool-use) mimarisi ile OpenRouter API aracılığıyla 120 Milyar Parametreli '
    'Açık Kaynak Büyük Dil Modelinden yararlanarak, sistem geleneksel düğme tabanlı chatbotların '
    'başaramadığı bir şeyi gerçekleştirir: doğal dil anlayışı tarafından yönlendirilen, gerçek anlamda '
    'otonom, çok adımlı görev yürütme. "Yarın sabah 10\'da bir kardiyologa randevu al" gibi tek bir '
    'hasta mesajı, 4\'ten fazla veritabanı işleminden oluşan bir zinciri tetikler - tamamı yapay zekâ '
    'tarafından herhangi bir sabit kodlanmış konuşma akışı olmadan orkestre edilir.'
)

doc.add_heading('Teknik Başarılar', level=2)
achievements = [
    '10 RESTful API rota grubu ve 9 arka uç denetleyicisi ile full-stack uygulama.',
    'Hem ön uçta (rota korumaları) hem de arka uçta (ara yazılım) uygulanan rol tabanlı erişim kontrolü.',
    'İşlem tabanlı çakışma algılama ile atomik randevu oluşturma.',
    'Profesyonel HTML şablonları kullanan 7 farklı e-posta türü ile otomatik e-posta bildirim sistemi.',
    '7 veritabanı aracı, kalıcı konuşma belleği ve çok adımlı araç zincirleme özellikli yapay zekâ chatbot.',
    'Zaman sınırlı tokenlarla e-posta tabanlı davet sistemi aracılığıyla güvenli personel kayıt süreci.',
    'Durum dağılımı, doktor bazlı istatistikler ve günlük trend analizi içeren yönetici analitik kontrol paneli.',
    'i18next kullanarak çoklu dil desteği altyapısı.',
]
for a in achievements:
    add_bullet(a)

doc.add_paragraph()

doc.add_paragraph(
    'Bu proje, büyük dil modellerinin pratik, alan odaklı uygulamalara entegrasyon potansiyelini '
    'sergiler - yapay zekâyı basit soru-cevaptan gerçek yazılım sistemleri içinde otonom görev '
    'yürütmeye taşır.'
)

# Alt Bilgi
doc.add_paragraph()
doc.add_paragraph()
line_para = doc.add_paragraph()
line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
line_run = line_para.add_run('_' * 60)
line_run.font.color.rgb = RGBColor(0x0d, 0x6e, 0xfd)
line_run.font.size = Pt(10)

footer_text = doc.add_paragraph()
footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_text.add_run('CarePlus - Online Hastane Yönetim Sistemi\nNisan 2026')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.italic = True


# ====== KAYDET ======
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Online_Hastane_Yonetim_Sistemi_Raporu_TR.docx')
doc.save(output_path)
print("Rapor kaydedildi: " + output_path)
