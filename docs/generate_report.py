"""
Generates a professional Word Document (.docx) report for the
Online Hospital Management System - suitable for academic presentation.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# Page Setup
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Style Configuration
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for level in range(1, 4):
    heading_style = doc.styles['Heading %d' % level]
    heading_style.font.color.rgb = RGBColor(0x0d, 0x47, 0xa1)
    heading_style.font.bold = True
    if level == 1:
        heading_style.font.size = Pt(20)
    elif level == 2:
        heading_style.font.size = Pt(15)
    else:
        heading_style.font.size = Pt(12)

# Code Block Style
code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = 'Consolas'
code_style.font.size = Pt(8.5)
code_style.font.color.rgb = RGBColor(0xD4, 0xD4, 0xD4)
code_style.paragraph_format.space_before = Pt(6)
code_style.paragraph_format.space_after = Pt(6)
code_style.paragraph_format.line_spacing = 1.15
code_style.paragraph_format.left_indent = Cm(0.5)


def add_code_block(code_text, caption=""):
    if caption:
        cap_para = doc.add_paragraph()
        cap_run = cap_para.add_run(caption)
        cap_run.bold = True
        cap_run.font.size = Pt(10)
        cap_run.font.color.rgb = RGBColor(0x0d, 0x47, 0xa1)
        cap_para.paragraph_format.space_after = Pt(2)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shading = parse_xml('<w:shd %s w:fill="1E1E1E"/>' % nsdecls("w"))
    cell._tc.get_or_add_tcPr().append(shading)
    cell.width = Inches(6.5)
    cell.paragraphs[0].clear()

    lines = code_text.strip().split('\n')
    for i, line in enumerate(lines):
        if i == 0:
            para = cell.paragraphs[0]
        else:
            para = cell.add_paragraph()
        para.style = doc.styles['CodeBlock']
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.1
        run = para.add_run(line if line else ' ')
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0xD4, 0xD4, 0xD4)
    doc.add_paragraph()


def add_bullet(text, bold_prefix=""):
    para = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_bold = para.add_run(bold_prefix)
        run_bold.bold = True
        para.add_run(text)
    else:
        para.add_run(text)
    return para


def add_key_value_table(data, col1_header="Component", col2_header="Technology"):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Shading Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = col1_header
    hdr[1].text = col2_header
    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    for key, value in data:
        row = table.add_row().cells
        row[0].text = key
        row[1].text = value
        for cell in row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


# ====== COVER PAGE ======
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('ONLINE HOSPITAL\nMANAGEMENT SYSTEM')
run.bold = True
run.font.size = Pt(32)
run.font.color.rgb = RGBColor(0x0d, 0x47, 0xa1)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('with AI-Powered Chatbot Integration')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x42, 0x42, 0x42)
run.italic = True

doc.add_paragraph()

line_para = doc.add_paragraph()
line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
line_run = line_para.add_run('_' * 60)
line_run.font.color.rgb = RGBColor(0x0d, 0x6e, 0xfd)
line_run.font.size = Pt(12)

doc.add_paragraph()

details = [
    ('Project Type:', 'Full-Stack Web Application'),
    ('Course:', 'Software Engineering'),
    ('Date:', 'April 2026'),
    ('Technologies:', 'React.js, Node.js, Express.js, MySQL, OpenAI API'),
]
for label, value in details:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label + ' ')
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    r2 = p.add_run(value)
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_page_break()

# ====== TABLE OF CONTENTS ======
doc.add_heading('Table of Contents', level=1)

toc_items = [
    ('1.', 'Introduction & Project Overview'),
    ('2.', 'System Architecture'),
    ('3.', 'Technology Stack'),
    ('4.', 'Database Design'),
    ('5.', 'Backend Implementation'),
    ('  5.1', 'Server Entry Point (index.js)'),
    ('  5.2', 'Database Configuration'),
    ('  5.3', 'JWT Authentication Middleware'),
    ('  5.4', 'Authentication Controller'),
    ('  5.5', 'Appointments Controller'),
    ('  5.6', 'Email Notification Service'),
    ('6.', 'Frontend Implementation'),
    ('  6.1', 'Application Routing (App.js)'),
    ('  6.2', 'API Service Layer'),
    ('7.', 'AI Chatbot Integration'),
    ('  7.1', 'Chatbot Controller - Backend'),
    ('  7.2', 'ChatAssistant - Frontend UI'),
    ('  7.3', 'Tool/Function Calling Architecture'),
    ('8.', 'Security Features'),
    ('9.', 'API Endpoints Reference'),
    ('10.', 'Conclusion'),
]

for num, item in toc_items:
    p = doc.add_paragraph()
    r1 = p.add_run(num + '  ')
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(item)
    r2.font.size = Pt(11)
    if num.startswith('  '):
        p.paragraph_format.left_indent = Cm(1.2)

doc.add_page_break()

# ====== 1. INTRODUCTION ======
doc.add_heading('1. Introduction & Project Overview', level=1)

doc.add_paragraph(
    'The Online Hospital Management System (CarePlus) is a comprehensive, full-stack web application '
    'designed to digitize and streamline the core operations of a modern hospital. The system provides '
    'role-based dashboards for four distinct user types - Administrators, Doctors, Patients, and '
    'Receptionists - each with tailored functionalities and interfaces.'
)

doc.add_paragraph(
    'A key differentiator of this project is the integration of an AI-powered chatbot that acts as a '
    'virtual receptionist. Powered by a 120-Billion Parameter Open Source Large Language Model (LLM) '
    'through the OpenRouter API, the chatbot can autonomously book, reschedule, view, and cancel '
    'appointments by directly interacting with the hospital\'s MySQL database through a secure '
    'function-calling (tool-use) architecture.'
)

doc.add_heading('Key Features', level=2)
features = [
    ('Multi-Role System: ', 'Patient, Doctor, Receptionist, and Admin dashboards with role-based access control (RBAC).'),
    ('Appointment Management: ', 'Full CRUD operations with conflict detection, status tracking, and email notifications.'),
    ('AI Chatbot: ', 'Natural language appointment management using a 120B parameter LLM with function calling.'),
    ('Email Notifications: ', 'Automated email notifications for booking confirmations, cancellations, rescheduling, and staff invitations.'),
    ('Staff Invitation System: ', 'Admin can invite doctors and receptionists via secure email-based onboarding links.'),
    ('Real-time Analytics: ', 'Admin dashboard with appointment statistics, per-doctor breakdown, and daily trend charts.'),
    ('Internationalization: ', 'Multi-language support using i18next framework.'),
    ('Security: ', 'JWT authentication, bcrypt password hashing, role-based route protection, and SQL injection prevention.'),
]
for prefix, text in features:
    add_bullet(text, prefix)

doc.add_page_break()

# ====== 2. SYSTEM ARCHITECTURE ======
doc.add_heading('2. System Architecture', level=1)

doc.add_paragraph(
    'The application follows a classic three-tier client-server architecture with a clear separation '
    'of concerns between the presentation layer (React), the business logic layer (Express/Node.js), '
    'and the data layer (MySQL). The AI chatbot adds a fourth intelligent layer that bridges the user '
    'interface with backend database operations autonomously.'
)

doc.add_heading('Architecture Overview', level=2)

arch_data = [
    ('Presentation Layer', 'React.js 18 SPA with Bootstrap UI - Port 3000'),
    ('API Layer', 'Express.js REST API with JWT auth middleware - Port 5000'),
    ('Business Logic', '9 Controllers: Auth, Appointments, Chatbot, Doctors, Clinics, Users, Messages, Notifications, Medical Records'),
    ('AI Layer', 'OpenRouter LLM API (120B OSS model) with 7 Function Tools'),
    ('Email Service', 'Nodemailer with Gmail SMTP - 7 email template types'),
    ('Data Layer', 'MySQL with Connection Pool (10 connections) - Parameterized Queries'),
]
add_key_value_table(arch_data, 'Layer', 'Technology & Details')

doc.add_heading('Project File Structure', level=2)
add_code_block('''online-hospital-system/
|-- client/                          # React Frontend
|   |-- src/
|   |   |-- components/
|   |   |   |-- admin/               # Admin panel components
|   |   |   |-- auth/                # Login & Register forms
|   |   |   |-- chat/                # AI ChatAssistant component
|   |   |   |-- common/              # Shared components
|   |   |   |-- doctor/              # Doctor panel components
|   |   |   |-- layout/              # Header, Footer, DashboardLayout
|   |   |   |-- patient/             # Patient panel components
|   |   |   |-- reception/           # Receptionist components
|   |   |-- pages/                   # Route-level page components
|   |   |-- services/
|   |   |   |-- api.js               # Axios HTTP client & domain services
|   |   |   |-- aiService.js         # Chatbot API service
|   |   |-- utils/                   # RBAC utilities (roleGuard.js)
|   |   |-- App.js                   # Main app with routing
|   |   |-- i18n.js                  # Internationalization config
|
|-- server/                          # Node.js Backend
|   |-- config/
|   |   |-- db.js                    # MySQL connection pool
|   |   |-- email.js                 # Nodemailer email service (7 templates)
|   |-- controllers/
|   |   |-- authController.js        # Registration, login, invitation system
|   |   |-- appointmentsController.js # CRUD + analytics + conflict detection
|   |   |-- chatbotController.js     # AI LLM integration with 7 tools
|   |   |-- clinicsController.js     # Department management
|   |   |-- doctorsController.js     # Doctor queries
|   |   |-- usersController.js       # User management (admin)
|   |   |-- messagesController.js    # Messaging system
|   |   |-- notificationsController.js # Notification management
|   |   |-- medicalRecordsController.js # Medical records
|   |-- middleware/
|   |   |-- auth.js                  # JWT verify + role-based authorize
|   |-- routes/                      # 10 Express route files
|   |-- seeder.js                    # Database seeder (clinics + 30 doctors)
|   |-- index.js                     # Express server entry point
|   |-- .env                         # Environment variables''', "Figure 2.1 - Complete Project File Structure")

doc.add_page_break()

# ====== 3. TECHNOLOGY STACK ======
doc.add_heading('3. Technology Stack', level=1)

doc.add_heading('Frontend Technologies', level=2)
add_key_value_table([
    ('UI Framework', 'React.js 18.2'),
    ('Routing', 'React Router DOM 6.14'),
    ('HTTP Client', 'Axios 1.4'),
    ('UI Components', 'React Bootstrap 2.8 + Bootstrap 5.3'),
    ('Charts', 'Chart.js 4.3 + react-chartjs-2'),
    ('Icons', 'React Icons 4.12'),
    ('Internationalization', 'i18next + react-i18next'),
    ('Date Handling', 'Moment.js 2.29'),
    ('Auth Token Parsing', 'jwt-decode 3.1'),
])

doc.add_heading('Backend Technologies', level=2)
add_key_value_table([
    ('Runtime', 'Node.js'),
    ('Web Framework', 'Express.js 5.1'),
    ('Database Driver', 'mysql2 3.14 (Promise-based connection pool)'),
    ('Authentication', 'JSON Web Tokens (jsonwebtoken 9.0)'),
    ('Password Hashing', 'bcrypt 6.0'),
    ('Email Service', 'Nodemailer 7.0 (Gmail SMTP)'),
    ('AI/LLM Integration', 'OpenAI SDK 6.34 -> OpenRouter API'),
    ('LLM Model', 'openai/gpt-oss-120b:free (120B param OSS model)'),
    ('Environment Config', 'dotenv 16.5'),
    ('Dev Server', 'Nodemon 3.1'),
])

doc.add_heading('Database', level=2)
add_key_value_table([
    ('RDBMS', 'MySQL'),
    ('Connection Strategy', 'Connection Pool (10 connections)'),
    ('ORM', 'Raw SQL with parameterized queries (SQL injection safe)'),
], 'Component', 'Details')

doc.add_page_break()

# ====== 4. DATABASE DESIGN ======
doc.add_heading('4. Database Design', level=1)

doc.add_paragraph(
    'The system uses a relational MySQL database with the following core tables. '
    'All tables are linked through foreign key relationships to maintain referential integrity.'
)

doc.add_heading('Entity Relationship Overview', level=2)
add_key_value_table([
    ('users', 'id, username, email, password, role, status, invitation_token, invitation_expires'),
    ('patients', 'id, user_id (FK->users), first_name, last_name, phone'),
    ('doctors', 'id, user_id (FK->users), first_name, last_name, specialization, qualification, phone, clinic_id (FK->clinics)'),
    ('clinics', 'id, name, description, status'),
    ('appointments', 'id, patient_id (FK->patients), doctor_id (FK->doctors), clinic_id (FK->clinics), appointment_date, appointment_time, reason, notes, status'),
    ('notifications', 'id, user_id (FK->users), message, type, related_id'),
    ('receptionists', 'id, user_id (FK->users), first_name, last_name'),
], 'Table', 'Key Columns')

doc.add_heading('Appointment Status Flow', level=2)
doc.add_paragraph('Appointments follow a well-defined lifecycle with the following valid states:')
statuses = [
    ('pending: ', 'Initial state when a patient books via form or chatbot.'),
    ('pending_assignment: ', 'Submitted without a specific doctor; awaiting receptionist assignment.'),
    ('confirmed: ', 'Doctor or receptionist has confirmed the appointment.'),
    ('scheduled: ', 'Appointment is officially on the calendar.'),
    ('in-progress: ', 'Patient is currently being seen.'),
    ('completed: ', 'Appointment has been completed.'),
    ('cancelled: ', 'Appointment was cancelled by patient, doctor, or admin.'),
    ('no-show: ', 'Patient did not attend the appointment.'),
]
for prefix, desc in statuses:
    add_bullet(desc, prefix)

doc.add_page_break()

# ====== 5. BACKEND IMPLEMENTATION ======
doc.add_heading('5. Backend Implementation', level=1)

doc.add_paragraph(
    'The backend is built with Express.js 5.1 running on Node.js. It follows an MVC-like architecture '
    'with controllers handling business logic, routes defining API endpoints, middleware managing '
    'authentication/authorization, and config modules for database and email services.'
)

# 5.1 Server Entry Point
doc.add_heading('5.1 Server Entry Point (index.js)', level=2)
doc.add_paragraph(
    'The main server file initializes Express, loads environment variables, registers middleware '
    '(CORS, JSON body parser), and wires all route modules to their respective API paths.'
)

add_code_block('''const express = require('express');
const cors = require('cors');
const dotenv = require('dotenv');
const path = require('path');

// Route imports
const authRoutes = require('./routes/auth');
const appointmentRoutes = require('./routes/appointments');
const usersRoutes = require('./routes/users');
const notificationsRoutes = require('./routes/notifications');
const messagesRoutes = require('./routes/messages');
const doctorsRoutes = require('./routes/doctors');
const clinicsRoutes = require('./routes/clinics');
const medicalRecordsRoutes = require('./routes/medical-records');
const patientsRoutes = require('./routes/patients');
const chatbotRoutes = require('./routes/chatbot');

dotenv.config({ path: path.resolve(__dirname, '.env') });
const app = express();
const PORT = process.env.PORT || 5000;

// Middleware
app.use(cors());
app.use(express.json());
app.use(express.urlencoded({ extended: true }));

// API Routes
app.use('/api/auth', authRoutes);
app.use('/api/appointments', appointmentRoutes);
app.use('/api/users', usersRoutes);
app.use('/api/notifications', notificationsRoutes);
app.use('/api/messages', messagesRoutes);
app.use('/api/doctors', doctorsRoutes);
app.use('/api/clinics', clinicsRoutes);
app.use('/api/medical-records', medicalRecordsRoutes);
app.use('/api/patients', patientsRoutes);
app.use('/api/chatbot', chatbotRoutes);

app.listen(PORT, () => {
  console.log(`Server running on port ${PORT}`);
});''', "Figure 5.1 - server/index.js: Express Server Initialization & Route Registration")

doc.add_paragraph(
    'The server exposes 10 RESTful API route groups, each mapped to a dedicated controller. '
    'CORS is enabled to allow the React frontend (port 3000) to communicate with the API (port 5000).'
)

# 5.2 Database Configuration
doc.add_heading('5.2 Database Configuration (config/db.js)', level=2)
doc.add_paragraph(
    'The database module creates a MySQL connection pool using mysql2 with Promise support. '
    'This allows efficient connection reuse and prevents connection exhaustion under load.'
)

add_code_block('''const mysql = require('mysql2');
const dotenv = require('dotenv');
const path = require('path');

dotenv.config({ path: path.join(__dirname, '../.env') });

const pool = mysql.createPool({
  host: process.env.DB_HOST || 'localhost',
  port: parseInt(process.env.DB_PORT) || 3306,
  user: process.env.DB_USER || 'root',
  password: process.env.DB_PASSWORD,
  database: process.env.DB_NAME || 'hospital_management',
  waitForConnections: true,
  connectionLimit: 10,
  queueLimit: 0
});

const promisePool = pool.promise();
module.exports = promisePool;''', "Figure 5.2 - server/config/db.js: MySQL Connection Pool with Promises")

doc.add_page_break()

# 5.3 Auth Middleware
doc.add_heading('5.3 JWT Authentication Middleware', level=2)
doc.add_paragraph(
    'All protected routes pass through a two-layer security middleware: '
    'authenticateToken verifies the JWT token from the Authorization header, '
    'and authorize enforces role-based access control.'
)

add_code_block('''const jwt = require('jsonwebtoken');

// Middleware to verify JWT token
const authenticateToken = (req, res, next) => {
  const authHeader = req.headers['authorization'];
  const token = authHeader && authHeader.split(' ')[1];
  
  if (!token) {
    return res.status(401).json({ message: 'Access denied. No token provided.' });
  }

  try {
    const decoded = jwt.verify(token, process.env.JWT_SECRET);
    req.user = decoded;
    next();
  } catch (error) {
    return res.status(403).json({ message: 'Invalid token.' });
  }
};

// Role-based authorization middleware
const authorize = (roles = []) => {
  if (typeof roles === 'string') {
    roles = [roles];
  }

  return (req, res, next) => {
    if (roles.length === 0) return next();
    if (!req.user || !roles.includes(req.user.role)) {
      return res.status(403).json({ 
        message: 'Forbidden: You do not have permission' 
      });
    }
    next();
  };
};

module.exports = { authenticateToken, authorize };''', "Figure 5.3 - server/middleware/auth.js: JWT Token Verification & Role Authorization")

# 5.4 Auth Controller
doc.add_heading('5.4 Authentication Controller (Highlights)', level=2)
doc.add_paragraph(
    'The authentication controller handles user registration, login, email verification, password management, '
    'and the staff invitation system. Below are the key functions:'
)

doc.add_heading('User Registration with Email Verification', level=3)
add_code_block('''exports.register = async (req, res) => {
  try {
    const { username, email, password, role, first_name, last_name, 
            phone, specialization, qualification, verificationCode } = req.body;

    // Email verification check
    if (process.env.EMAIL_VERIFICATION_ENABLED === 'true') {
      if (!verificationCodes[email] || !verificationCodes[email].verified) {
        // Send verification code via email
      }
    }

    // Role restriction: Only admins can register non-patient roles
    if (role !== 'patient') {
      const authHeader = req.headers['authorization'];
      const token = authHeader && authHeader.split(' ')[1];
      const decoded = jwt.verify(token, process.env.JWT_SECRET);
      if (decoded.role !== 'admin') 
        return res.status(403).json({ message: 'Admin privileges required.' });
    }

    // Password hashing with bcrypt (10 salt rounds)
    const salt = await bcrypt.genSalt(10);
    const hashedPassword = await bcrypt.hash(password, salt);

    // Transactional insert: user + role-specific profile
    await db.query('START TRANSACTION');
    const [result] = await db.query(
      'INSERT INTO users (username, email, password, role) VALUES (?, ?, ?, ?)',
      [username, email, hashedPassword, role]
    );

    if (role === 'doctor') {
      await db.query(
        'INSERT INTO doctors (user_id, first_name, last_name, specialization) VALUES (...)',
        [result.insertId, first_name, last_name, specialization]
      );
    } else if (role === 'patient') {
      await db.query(
        'INSERT INTO patients (user_id, first_name, last_name, phone) VALUES (...)',
        [result.insertId, first_name, last_name, phone]
      );
    }

    const token = jwt.sign(
      { id: result.insertId, username, email, role },
      process.env.JWT_SECRET, { expiresIn: '1d' }
    );

    await db.query('COMMIT');
    res.status(201).json({ message: 'User registered successfully', token });
  } catch (error) {
    await db.query('ROLLBACK'); throw error;
  }
};''', "Figure 5.4a - authController.js: User Registration with Transaction & Role Handling")

doc.add_page_break()

doc.add_heading('User Login with JWT Generation', level=3)
add_code_block('''exports.login = async (req, res) => {
  try {
    const { username, password } = req.body;

    const [users] = await db.query(
      'SELECT * FROM users WHERE username = ? OR email = ?',
      [username, username]  // Allow login with username or email
    );

    if (users.length === 0) 
      return res.status(401).json({ message: 'Invalid credentials' });

    const user = users[0];
    const isPasswordValid = await bcrypt.compare(password, user.password);
    if (!isPasswordValid) 
      return res.status(401).json({ message: 'Invalid credentials' });

    // Generate JWT token with user data (expires in 24h)
    const token = jwt.sign(
      { id: user.id, username: user.username, email: user.email, role: user.role },
      process.env.JWT_SECRET || 'hospital_system_jwt_secret_key',
      { expiresIn: '1d' }
    );

    res.json({ 
      token, 
      user: { id: user.id, username: user.username, role: user.role } 
    });
  } catch (error) {
    res.status(500).json({ message: 'Server error' });
  }
};''', "Figure 5.4b - authController.js: Login with bcrypt Password Comparison & JWT Token")

# 5.5 Appointments Controller
doc.add_heading('5.5 Appointments Controller (Highlights)', level=2)
doc.add_paragraph(
    'The appointments controller is the largest backend module (~570 lines). It handles creation with '
    'atomic conflict checking, status updates with cascading email notifications, analytics, and '
    'ownership verification.'
)

add_code_block('''exports.createAppointment = async (req, res) => {
  try {
    const { doctorId, clinicId, date, time, reason, notes, patientUserId } = req.body;
    const isReceptionist = req.user.role === 'receptionist';
    const targetUserId = isReceptionist && patientUserId ? patientUserId : req.user.id;

    // Resolve Patient Profile
    const [patients] = await db.query(
      'SELECT id, first_name, last_name FROM patients WHERE user_id = ?', 
      [targetUserId]
    );
    if (patients.length === 0) 
      return res.status(404).json({ message: 'Patient record not found.' });

    // Atomic Conflict Check & Insert (Transaction)
    const conn = await db.getConnection();
    await conn.beginTransaction();

    if (doctorId) {
      const [conflict] = await conn.query(
        `SELECT id FROM appointments
         WHERE doctor_id = ? AND appointment_date = ?
           AND TIME_FORMAT(appointment_time,'%H:%i') = TIME_FORMAT(?,'%H:%i')
           AND status NOT IN ('cancelled','no-show')`,
        [doctorId, date, formattedTime]
      );
      if (conflict.length > 0) {
        await conn.rollback(); conn.release();
        return res.status(409).json({ message: 'Time slot already booked.' });
      }
    }

    const [result] = await conn.query(
      'INSERT INTO appointments (...) VALUES (?, ?, ?, ?, ?, ?, ?, ?)',
      [patient.id, doctorId, clinicId, date, formattedTime, reason, notes, 'pending']
    );

    await conn.commit(); conn.release();

    // Non-blocking email notifications to doctor
    emailService.sendDoctorNewAppointmentEmail(...).catch(console.error);

    return res.status(201).json({ 
      message: 'Appointment created successfully', 
      id: result.insertId 
    });
  } catch (error) { ... }
};''', "Figure 5.5 - appointmentsController.js: Atomic Appointment Creation with Conflict Detection")

doc.add_page_break()

# 5.6 Email Service
doc.add_heading('5.6 Email Notification Service', level=2)
doc.add_paragraph(
    'The system uses Nodemailer with Gmail SMTP to send automated email notifications. '
    'The email service provides 7 distinct email types, each with professional HTML templates:'
)

email_types = [
    'Email Verification (6-digit OTP for registration)',
    'Staff Invitation (with secure setup link, 48-hour expiry)',
    'Appointment Pending (acknowledgement to patient)',
    'Appointment Confirmation (with doctor, date, time, clinic details)',
    'Appointment Rescheduled (updated details notification)',
    'Appointment Cancellation (with cancellation reason)',
    'Doctor New Appointment (notification to assigned doctor)',
]
for t in email_types:
    add_bullet(t)

add_code_block('''const nodemailer = require('nodemailer');

const transporter = nodemailer.createTransport({
  service: 'gmail',
  auth: {
    user: process.env.SMTP_USER,
    pass: process.env.SMTP_PASS  // Gmail App Password
  }
});

const sendAppointmentConfirmationEmail = async (
  patientEmail, patientName, doctorName, date, time, clinic
) => {
  const mailOptions = {
    from: SMTP_USER,
    to: patientEmail,
    subject: 'Appointment Confirmed - Hospital Management System',
    html: `
      <div style="font-family: Arial; max-width: 600px; margin: auto;">
        <h2>Appointment Confirmation</h2>
        <p>Dear ${patientName},</p>
        <p>Your appointment has been confirmed:</p>
        <ul>
          <li><strong>Doctor:</strong> ${doctorName}</li>
          <li><strong>Date:</strong> ${date}</li>
          <li><strong>Time:</strong> ${time}</li>
          <li><strong>Clinic:</strong> ${clinic}</li>
        </ul>
        <p>Please arrive 15 minutes before your appointment.</p>
      </div>`
  };
  await transporter.sendMail(mailOptions);
};''', "Figure 5.6 - config/email.js: Nodemailer SMTP Setup & Email Template Example")

doc.add_page_break()

# ====== 6. FRONTEND IMPLEMENTATION ======
doc.add_heading('6. Frontend Implementation', level=1)

doc.add_paragraph(
    'The frontend is a React.js 18 single-page application (SPA) using React Router for '
    'client-side navigation, React Bootstrap for UI components, and Axios for API communication. '
    'The application uses a component-based architecture with shared layouts and role-specific pages.'
)

# 6.1 App.js
doc.add_heading('6.1 Application Routing & Role-Based Access (App.js)', level=2)
doc.add_paragraph(
    'The main App component implements a complete role-based routing system. A ProtectedRoute wrapper '
    'component enforces both authentication (redirect to /login) and authorization (403 screen for '
    'unauthorized roles). Each role has its own route group and dedicated dashboard layout.'
)

add_code_block('''function App() {
  const [isAuthenticated, setIsAuthenticated] = useState(false);
  const [user, setUser] = useState(null);

  // ProtectedRoute - enforces authentication + role authorization
  const ProtectedRoute = ({ children, allowedRoles = [] }) => {
    if (!isAuthenticated) return <Navigate to="/login" replace />;
    if (allowedRoles.length > 0 && !allowedRoles.includes(user?.role)) 
      return <Unauthorized />;
    return children;
  };

  return (
    <Router>
      <DashboardLayout user={user} logout={logout}>
        <Routes>
          {/* Admin Routes */}
          <Route path="/admin" element={
            <ProtectedRoute allowedRoles={['admin']}>
              <AdminDashboard />
            </ProtectedRoute>
          } />
          <Route path="/admin/doctors" element={...} />
          <Route path="/admin/appointments" element={...} />
          <Route path="/admin/users" element={...} />
          <Route path="/admin/departments" element={...} />

          {/* Doctor Routes */}
          <Route path="/doctor" element={
            <ProtectedRoute allowedRoles={['doctor']}>
              <DoctorDashboard />
            </ProtectedRoute>
          } />
          <Route path="/doctor/appointments" element={...} />

          {/* Patient Routes */}
          <Route path="/patient" element={...} />
          <Route path="/patient/appointments/book" element={...} />

          {/* Reception Routes */}
          <Route path="/reception" element={...} />
          <Route path="/reception/appointments" element={...} />

          {/* Smart Home Redirect */}
          <Route path="/" element={
            <Navigate to={getHomePath(user?.role)} replace />
          } />
        </Routes>
      </DashboardLayout>
    </Router>
  );
}''', "Figure 6.1 - client/src/App.js: Role-Based Protected Routing System")

doc.add_page_break()

# 6.2 API Service
doc.add_heading('6.2 Centralized API Service Layer', level=2)
doc.add_paragraph(
    'All HTTP communication is centralized through an Axios-based API service module. '
    'Request interceptors automatically attach JWT tokens, and response interceptors handle '
    'global error cases like session expiry (401 -> auto-redirect to login).'
)

add_code_block('''import axios from 'axios';

const API_BASE_URL = process.env.REACT_APP_API_URL || 'http://localhost:5000/api';

const api = axios.create({
  baseURL: API_BASE_URL,
  headers: { 'Content-Type': 'application/json' }
});

// Auto-attach JWT token to every outgoing request
api.interceptors.request.use((config) => {
  const token = localStorage.getItem('token');
  if (token) config.headers.Authorization = `Bearer ${token}`;
  return config;
});

// Auto-handle expired sessions (401 -> redirect to login)
api.interceptors.response.use(
  (response) => response,
  (error) => {
    if (error.response?.status === 401) {
      localStorage.removeItem('token');
      localStorage.removeItem('user');
      window.location.href = '/login';
    }
    return Promise.reject(error);
  }
);

// Domain Service Objects
export const appointmentService = {
  getAll:         () => api.get('/appointments'),
  create:        (data) => api.post('/appointments', data),
  update:    (id, data) => api.patch(`/appointments/${id}/status`, data),
  getAnalytics:  () => api.get('/appointments/analytics'),
};

export const chatbotService = {
  sendMessage: (message) => api.post('/chatbot/message', { message }),
};''', "Figure 6.2 - client/src/services/api.js: Axios Interceptors & Domain Services")

doc.add_page_break()

# ====== 7. AI CHATBOT INTEGRATION ======
doc.add_heading('7. AI Chatbot Integration', level=1)

doc.add_paragraph(
    'The most innovative feature of this system is the AI-powered chatbot that serves as a virtual '
    'hospital receptionist. Unlike traditional rule-based chatbots, this implementation uses a '
    '120-Billion Parameter Open Source Large Language Model (LLM) through the OpenRouter API, '
    'enabling natural language understanding and autonomous task execution.'
)

doc.add_heading('How It Works', level=2)
doc.add_paragraph(
    'The chatbot architecture is based on a "Tool Use / Function Calling" pattern where the LLM '
    'does not just generate text - it can actively request the server to execute database operations. '
    'The flow works as follows:'
)

steps = [
    'Patient sends a natural language message (e.g., "I want to see a cardiologist tomorrow").',
    'The message is sent to the backend chatbot controller along with conversation history.',
    'The controller forwards the conversation to the OpenRouter LLM API with defined tools.',
    'The LLM analyzes the request and outputs a structured JSON tool call (e.g., get_clinics).',
    'The server executes the database query and sends the result back to the LLM.',
    'The LLM may chain multiple tool calls (get_clinics -> get_doctors -> check_availability -> book).',
    'Once all data is gathered, the LLM generates a final natural language response to the patient.',
]
for i, step in enumerate(steps, 1):
    add_bullet(step, 'Step %d: ' % i)

doc.add_page_break()

# 7.1 Chatbot Controller
doc.add_heading('7.1 Chatbot Controller - Backend (chatbotController.js)', level=2)

doc.add_heading('LLM Initialization & System Prompt', level=3)
doc.add_paragraph(
    'The controller initializes the OpenAI SDK pointing to the OpenRouter API endpoint. '
    'A comprehensive system instruction defines the AI\'s persona and operational rules.'
)

add_code_block('''const OpenAI = require('openai');
const db = require('../config/db');

// Initialize OpenRouter Client (Using OpenAI SDK with custom base URL)
const openai = new OpenAI({
  baseURL: 'https://openrouter.ai/api/v1',
  apiKey: process.env.OPENROUTER_API_KEY || 'MISSING_KEY'
});

const systemInstruction = `You are CarePlus Assistant, a helpful AI 
virtual receptionist for a hospital. 
Your primary job is to help patients book, reschedule, view, and cancel 
their medical appointments.
You have access to backend functional tools to interact with the database.

CRITICAL RULES FOR BOOKING:
1. When booking, you must collect the Clinic, Doctor, Date, and Time.
2. Call 'get_clinics' to show departments, then 'get_doctors'.
3. You MUST call 'check_availability' before finalizing a booking.
4. After verifying availability, ask the user to confirm.

CRITICAL RULES FOR RESCHEDULING:
1. Always call 'get_my_appointments' first to see existing bookings.
2. You MUST call 'check_availability' for the doctor on the new date.
3. Only call 'reschedule_appointment' if the time is available.`;''', "Figure 7.1a - chatbotController.js: OpenRouter LLM Client & System Instruction")

doc.add_heading('Tool/Function Definitions', level=3)
doc.add_paragraph(
    'The LLM is provided with a structured array of 7 tools it can invoke. Each tool has '
    'a name, description, and a JSON schema for its parameters:'
)

add_code_block('''const tools = [
  {
    type: 'function',
    function: {
      name: 'get_clinics',
      description: 'Get a list of all hospital clinics/departments.',
      parameters: { type: 'object', properties: {} } 
    }
  },
  {
    type: 'function',
    function: {
      name: 'get_doctors',
      description: 'Get doctors for a specific clinic.',
      parameters: {
        type: 'object',
        properties: {
          clinicId: { type: 'string', description: 'The clinic ID.' }
        },
        required: ['clinicId']
      }
    }
  },
  {
    type: 'function',
    function: {
      name: 'check_availability',
      description: 'Look up available timeslots for a doctor on a date.',
      parameters: {
        type: 'object',
        properties: {
          doctorId: { type: 'string' },
          date: { type: 'string', description: 'YYYY-MM-DD format.' }
        },
        required: ['doctorId', 'date']
      }
    }
  },
  {
    type: 'function',
    function: {
      name: 'book_appointment',
      description: 'Book a new appointment.',
      parameters: {
        type: 'object',
        properties: {
          clinicId: { type: 'string' }, doctorId: { type: 'string' },
          date: { type: 'string' },     time: { type: 'string' }
        },
        required: ['clinicId', 'doctorId', 'date', 'time']
      }
    }
  },
  // ... get_my_appointments, cancel_appointment, reschedule_appointment
];''', "Figure 7.1b - chatbotController.js: LLM Tool Definitions (JSON Schema)")

doc.add_page_break()

doc.add_heading('Message Processing Loop with Multi-Step Tool Chaining', level=3)
doc.add_paragraph(
    'The core of the chatbot is a while-loop that keeps sending tool results back to the LLM '
    'until it produces a final text response. This allows the AI to chain multiple database '
    'operations autonomously within a single user request.'
)

add_code_block('''// In-memory session store for conversational memory
const activeChats = {};

exports.processMessage = async (req, res) => {
  const userId = req.user.id;
  const { message } = req.body;

  // Initialize conversation history for new users
  if (!activeChats[userId]) {
    activeChats[userId] = [
      { role: 'system', content: systemInstruction }
    ];
  }
  
  activeChats[userId].push({ role: 'user', content: message });

  let isMakingToolCalls = true;
  let finalContent = "";

  // Processing loop: auto-handle multiple tool-call roundtrips
  while (isMakingToolCalls) {
    const response = await openai.chat.completions.create({
      model: 'openai/gpt-oss-120b:free', 
      messages: activeChats[userId],
      tools: tools,
    });

    const responseMessage = response.choices[0].message;
    activeChats[userId].push(responseMessage);

    if (responseMessage.tool_calls && responseMessage.tool_calls.length > 0) {
      // Execute each tool and feed results back to LLM
      for (const call of responseMessage.tool_calls) {
        const args = JSON.parse(call.function.arguments);
        const toolResult = await executeTool(
          call.function.name, args, userId
        );
        
        activeChats[userId].push({
          tool_call_id: call.id,
          role: 'tool',
          name: call.function.name,
          content: JSON.stringify(toolResult)
        });
      }
      // Loop continues - LLM will process tool results
    } else {
      // No more tool calls - final text response ready
      isMakingToolCalls = false;
      finalContent = responseMessage.content;
    }
  }
  
  return res.json({ text: finalContent });
};''', "Figure 7.1c - chatbotController.js: Multi-Step Tool Calling Loop with Memory")

doc.add_page_break()

doc.add_heading('Tool Executor Functions', level=3)
doc.add_paragraph(
    'Each tool maps to a specific database operation. The LLM never writes SQL directly - '
    'it outputs structured JSON parameters, and the server safely executes parameterized queries.'
)

add_code_block('''async function executeTool(name, args, userId) {
  switch(name) {
    case 'get_clinics': {
      const [clinics] = await db.query('SELECT id, name FROM clinics');
      return { clinics };
    }
    
    case 'get_doctors': {
      const [doctors] = await db.query(
        'SELECT id, first_name, last_name, specialization ' +
        'FROM doctors WHERE clinic_id = ?',
        [args.clinicId]
      );
      return { doctors };
    }
    
    case 'check_availability': {
      const [booked] = await db.query(
        `SELECT TIME_FORMAT(appointment_time,'%H:%i') as t 
         FROM appointments 
         WHERE doctor_id = ? AND appointment_date = ? 
           AND status NOT IN ('cancelled')`,
        [args.doctorId, args.date]
      );
      const bookedTimes = booked.map(r => r.t);
      // Generate all 30-min slots from 9 AM to 5 PM
      const allSlots = [];
      for (let h = 9; h < 17; h++) {
        allSlots.push(`${String(h).padStart(2,'0')}:00`);
        allSlots.push(`${String(h).padStart(2,'0')}:30`);
      }
      const available = allSlots.filter(s => !bookedTimes.includes(s));
      return { available_slots: available };
    }
    
    case 'book_appointment': {
      const [patients] = await db.query(
        'SELECT id FROM patients WHERE user_id = ?', [userId]
      );
      // Double-check conflict before INSERT
      const [conflict] = await db.query(
        "SELECT id FROM appointments WHERE doctor_id = ? " +
        "AND appointment_date = ? AND status NOT IN ('cancelled')",
        [args.doctorId, args.date]
      );
      if (conflict.length > 0) 
        return { status: 'error', message: 'Slot taken.' };

      await db.query(
        'INSERT INTO appointments (...) VALUES (?, ?, ?, ?, ?, ?, ?)',
        [patients[0].id, args.doctorId, args.clinicId, 
         args.date, args.time, 'Chatbot AI Booking', 'pending']
      );
      return { status: 'success', message: 'Appointment booked.' };
    }
    
    case 'cancel_appointment': {
      await db.query(
        "UPDATE appointments SET status = 'cancelled' WHERE id = ?",
        [args.appointmentId]
      );
      return { status: 'success' };
    }
    
    case 'reschedule_appointment': {
      // Verify new slot is free, then update
      const [conflict] = await db.query(...);
      if (conflict.length > 0) 
        return { status: 'error', message: 'Slot unavailable.' };
      await db.query(
        "UPDATE appointments SET appointment_date = ?, " +
        "appointment_time = ?, status = 'pending' WHERE id = ?",
        [args.newDate, args.newTime, args.appointmentId]
      );
      return { status: 'success', message: 'Rescheduled.' };
    }
  }
}''', "Figure 7.1d - chatbotController.js: Tool Executor - Secure Database Operations")

doc.add_page_break()

# 7.2 ChatAssistant Frontend
doc.add_heading('7.2 ChatAssistant - Frontend UI Component', level=2)
doc.add_paragraph(
    'The ChatAssistant is a floating React component that appears on all patient pages. '
    'It features a message bubble UI, typing indicators, markdown rendering, and persistent '
    'message history via localStorage.'
)

add_code_block('''import React, { useState, useEffect } from 'react';
import { Card, Button, Form, InputGroup } from 'react-bootstrap';
import { FaRobot, FaTimes, FaCommentMedical, FaPaperPlane } from 'react-icons/fa';
import { processChatMessage } from '../../services/aiService';

const ChatAssistant = () => {
  const [isOpen, setIsOpen] = useState(false);
  const [input, setInput] = useState('');
  const [isTyping, setIsTyping] = useState(false);
  
  // Persistent message history via localStorage
  const [messages, setMessages] = useState(() => {
    const saved = localStorage.getItem('chat_messages');
    return saved ? JSON.parse(saved) : [
      { sender: 'bot', 
        text: "Hi! I'm your CarePlus AI Assistant." }
    ];
  });

  // Sync messages to localStorage on every change
  useEffect(() => {
    localStorage.setItem('chat_messages', JSON.stringify(messages));
  }, [messages]);

  const handleSend = async () => {
    const displayMsg = input.trim();
    if (!displayMsg) return;

    // Show user message immediately
    setMessages(prev => [...prev, { sender: 'user', text: displayMsg }]);
    setInput('');
    setIsTyping(true);

    try {
      const response = await processChatMessage(null, displayMsg, null);
      setTimeout(() => {
        setIsTyping(false);
        setMessages(prev => [...prev, { 
          sender: 'bot', text: response.text 
        }]);
      }, 500);  // Small delay for natural feel
    } catch (error) {
      setIsTyping(false);
      setMessages(prev => [...prev, { 
        sender: 'bot', text: 'Error. Please try again.' 
      }]);
    }
  };

  // Markdown parser for bold text and bullet points
  const parseMarkdown = (text) => {
    if (!text) return '';
    let html = text.replace(/\\*\\*(.*?)\\*\\*/g, '<b>$1</b>');
    html = html.replace(/^\\* (.*$)/gim, '  $1');
    return html;
  };

  return (
    <>
      {/* Floating trigger button (bottom-right corner) */}
      {!isOpen && (
        <Button variant="primary" 
                className="rounded-circle shadow-lg position-fixed"
                style={{ width: 60, height: 60, bottom: 30, right: 30 }}
                onClick={() => setIsOpen(true)}>
          <FaCommentMedical size={26} />
        </Button>
      )}

      {/* Chat window: header + messages + input */}
      {isOpen && (
        <Card className="position-fixed shadow" 
              style={{ width: 380, height: 600, 
                       bottom: 30, right: 30, zIndex: 1050 }}>
          {/* Header with bot icon and title */}
          {/* Messages area with auto-scroll */}
          {/* Typing indicator (3 bouncing dots) */}
          {/* Input area with send button */}
        </Card>
      )}
    </>
  );
};''', "Figure 7.2 - ChatAssistant.js: Floating AI Chat UI Component")

doc.add_page_break()

# 7.3 Tool Architecture Summary
doc.add_heading('7.3 Tool/Function Calling Architecture Summary', level=2)

add_key_value_table([
    ('get_clinics', 'Reads all hospital departments from the clinics table.'),
    ('get_doctors', 'Filters doctors by clinic_id; returns name and specialization.'),
    ('check_availability', 'Generates 30-min slots (9AM-5PM), filters out booked ones.'),
    ('book_appointment', 'Double-checks conflicts, then INSERTs a new pending appointment.'),
    ('get_my_appointments', 'JOINs appointments with doctors and clinics for the logged-in patient.'),
    ('cancel_appointment', 'Sets appointment status to "cancelled" by ID.'),
    ('reschedule_appointment', 'Double-checks new slot, then UPDATEs date/time.'),
], 'Tool Name', 'Database Operation')

doc.add_page_break()

# ====== 8. SECURITY FEATURES ======
doc.add_heading('8. Security Features', level=1)

security_items = [
    ('JWT Authentication: ', 'All API endpoints (except login/register) require a valid JWT token in the Authorization header. Tokens expire after 24 hours.'),
    ('bcrypt Password Hashing: ', 'User passwords are salted and hashed with bcrypt (10 rounds) before storage. Plain-text passwords are never stored.'),
    ('Role-Based Access Control: ', 'Both frontend routes and backend endpoints enforce role checks. Admin-only operations are protected at both layers.'),
    ('SQL Injection Prevention: ', 'All database queries use parameterized placeholders (?). The AI chatbot never generates SQL - it outputs structured JSON parameters.'),
    ('Transactional Integrity: ', 'Critical operations use MySQL transactions (START TRANSACTION / COMMIT / ROLLBACK) to prevent partial writes.'),
    ('Staff Invitation Security: ', 'Staff onboarding uses cryptographically random tokens (crypto.randomBytes) with 48-hour expiration.'),
    ('Rate Limit Handling: ', 'The chatbot controller gracefully handles HTTP 429 responses from the AI provider, returning a user-friendly message.'),
    ('Session Isolation: ', 'Each patient\'s chatbot conversation is stored in an isolated in-memory session (activeChats[userId]).'),
]

for prefix, text in security_items:
    add_bullet(text, prefix)

doc.add_page_break()

# ====== 9. API ENDPOINTS REFERENCE ======
doc.add_heading('9. API Endpoints Reference', level=1)

doc.add_heading('Authentication', level=2)
add_key_value_table([
    ('POST /api/auth/register', 'Register new user (patient: public, others: admin only)'),
    ('POST /api/auth/login', 'Login and receive JWT token'),
    ('GET /api/auth/me', 'Get current user profile'),
    ('PATCH /api/auth/change-password', 'Update password (requires current password)'),
    ('POST /api/auth/invite', 'Invite staff member via email (admin only)'),
    ('GET /api/auth/verify-invite', 'Validate invitation token'),
    ('POST /api/auth/setup-invited', 'Complete invited account setup'),
], 'Endpoint', 'Description')

doc.add_heading('Appointments', level=2)
add_key_value_table([
    ('GET /api/appointments', 'Get appointments (auto-filtered by user role)'),
    ('POST /api/appointments', 'Create new appointment'),
    ('PATCH /api/appointments/:id/status', 'Update status / reschedule / assign doctor'),
    ('DELETE /api/appointments/:id', 'Delete appointment'),
    ('GET /api/appointments/analytics', 'Get dashboard analytics (admin)'),
    ('GET /api/appointments/availability', 'Check time slot availability'),
], 'Endpoint', 'Description')

doc.add_heading('Doctors, Clinics & Patients', level=2)
add_key_value_table([
    ('GET /api/doctors', 'Get all doctors'),
    ('GET /api/doctors/clinic/:clinicId', 'Get doctors by department'),
    ('GET /api/clinics', 'Get all departments'),
    ('POST /api/clinics', 'Create department (admin)'),
    ('PUT /api/clinics/:id', 'Update department (admin)'),
    ('DELETE /api/clinics/:id', 'Delete department (admin)'),
    ('GET /api/users', 'List all users (admin)'),
    ('GET /api/users/stats', 'Dashboard statistics (admin)'),
], 'Endpoint', 'Description')

doc.add_heading('AI Chatbot', level=2)
add_key_value_table([
    ('POST /api/chatbot/message', 'Send message to AI assistant (authenticated patients)'),
], 'Endpoint', 'Description')

doc.add_page_break()

# ====== 10. CONCLUSION ======
doc.add_heading('10. Conclusion', level=1)

doc.add_paragraph(
    'The Online Hospital Management System (CarePlus) demonstrates a comprehensive, production-grade '
    'approach to healthcare information systems. Built with modern web technologies - React.js on the '
    'frontend, Node.js/Express on the backend, and MySQL for data persistence - the system successfully '
    'implements all core hospital management functionalities while maintaining clean code architecture '
    'and robust security practices.'
)

doc.add_paragraph(
    'The AI chatbot integration represents the most technically ambitious component of the project. '
    'By leveraging a 120-Billion Parameter Open Source LLM through the OpenRouter API with a '
    'function-calling (tool-use) architecture, the system achieves something that traditional '
    'button-based chatbots cannot: truly autonomous, multi-step task execution driven by natural '
    'language understanding. A single patient message like "book me an appointment with a cardiologist '
    'tomorrow at 10 AM" triggers a chain of 4+ database operations - all orchestrated by the AI '
    'without any hardcoded conversation flows.'
)

doc.add_heading('Technical Achievements', level=2)
achievements = [
    'Full-stack application with 10 RESTful API route groups and 9 backend controllers.',
    'Role-based access control implemented at both frontend (route guards) and backend (middleware) layers.',
    'Atomic appointment booking with transaction-based conflict detection.',
    'Automated email notification system with 7 distinct email types using professional HTML templates.',
    'AI chatbot with 7 database tools, persistent conversation memory, and multi-step tool chaining.',
    'Secure staff onboarding through email-based invitation system with time-limited tokens.',
    'Admin analytics dashboard with status breakdown, per-doctor statistics, and daily trend analysis.',
    'Multi-language support infrastructure using i18next.',
]
for a in achievements:
    add_bullet(a)

doc.add_paragraph()

doc.add_paragraph(
    'This project showcases the potential of integrating large language models into practical, '
    'domain-specific applications - moving AI from simple question-answering to autonomous task '
    'execution within real software systems.'
)

# Footer
doc.add_paragraph()
doc.add_paragraph()
line_para = doc.add_paragraph()
line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
line_run = line_para.add_run('_' * 60)
line_run.font.color.rgb = RGBColor(0x0d, 0x6e, 0xfd)
line_run.font.size = Pt(10)

footer_text = doc.add_paragraph()
footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_text.add_run('CarePlus - Online Hospital Management System\nApril 2026')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.italic = True


# ====== SAVE ======
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Online_Hospital_Management_System_Report.docx')
doc.save(output_path)
print("Report saved to: " + output_path)
